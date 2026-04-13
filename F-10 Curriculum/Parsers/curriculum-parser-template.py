"""
Victorian Curriculum V2.0 - Generic Subject Parser
===================================================

VERSION HISTORY:
---------------
v4.5 - February 1, 2026 (EAL SUPPORT)
  - NEW: English as an Additional Language (EAL) support
  - EAL uses pathway levels (AL, A1, A2, BL, B1-B3, CL, C1-C4) not year levels
  - V2_CODE_PATTERN updated: VC2EAL[ABC][L1234][LRW]\d+ format
  - transform_band_for_output(): EAL levels stay as "Level AL" etc (not "Year")
  - transform_band_for_output(): Arts "Levels X and Y" -> "Years X and Y"
  - _generate_asc_code(): EAL format VC2EALALASC01, VC2EALA1ASC01, etc.
  - NEW: get_full_paragraph_text() extracts text from nested XML w:t elements
    (fixes VCAA docs where para.text returns empty but content exists in XML)
  - AS parsing now tracks strand headings within AS section for accurate strand assignment
  - CSV headers use ContentCode/ContentDescription (not _V2 suffix) for consistency

v4.4 - January 31, 2026 (BAND NAME TRANSFORMATION)
  - CRITICAL FIX: Added transform_band_for_output() function
  - Transforms document headings to standard Australian school terminology:
    * 'Foundation Level A' -> 'Foundation A'
    * 'Foundation Level B' -> 'Foundation B'
    * 'Foundation Level C' -> 'Foundation C'
    * 'Foundation Level D' -> 'Foundation D'
    * 'Level 1' -> 'Year 1' (through to Level 10 -> Year 10)
    * 'Level 10A' -> 'Year 10A'
  - Aligns with ACARA parser output format
  - Matches database schema and Australian school terminology

v4.3 - January 31, 2026 (MATHEMATICS FOUNDATION A-D & 10A FIX)
  - CRITICAL FIX: V2_CODE_PATTERN regex now captures all code formats:
    * Foundation pathway codes: VC2MFAN01, VC2MFBN01, VC2MFCN01, VC2MFDN01
    * Level 10A codes: VC2M10AN01, VC2M10AA01, etc.
    * Standard codes: VC2MFN01, VC2M7N01, VC2M10A01
    * Arts codes: VC2ADAFE01, VC2ADA6C01, VC2ADA10P03
  - Updated regex from: r'VC2[A-Z]+[F\d]+[A-Z]\d+'
    to: r'VC2[A-Z]+(?:F[A-D]?|\d{1,2})[A-Z]+\d+'
  - Mathematics now correctly parses 327 CDs across 16 bands
  - Supports Foundation Level A, B, C, D as separate bands
  - Supports Level 10A as separate band

v4.2 - January 30, 2026 (ALIGNED VERSION)
  - ALIGNMENT: Added ASComponentCode generation (VC2{prefix}{band}ASC{seq})
  - ALIGNMENT: AS Components CSV uses ASComponentText column name (not Text)
  - ALIGNMENT: Column order standardised to match ACARA parser v4.2
  - All CSV formats now fully aligned between Victorian and ACARA parsers

v4.0 - January 30, 2026
  - ALIGNMENT: Added SubjectArea and Subject columns to all CSV outputs
  - ALIGNMENT: Standardised CSV column order to match ACARA parser
  - ALIGNMENT: Added three-phase test infrastructure (Pre-check, Parse, Post-check)
  - ALIGNMENT: Standardised keyword extraction (min length 3, no hard limit)
  - ALIGNMENT: Added capitalisation validation in post-parse tests
  - ALIGNMENT: Improved stop words list (merged best from both parsers)
  - ENHANCEMENT: Added dataclass definitions for type safety
  - ENHANCEMENT: Pre-check validation includes nested table detection
  - Maintains all v3.1 features (nested tables, two-pass parsing, etc.)

v3.1 - December 29, 2024 (FINAL FIXED v2)
  - CRITICAL FIX: clean_text() now preserves newlines

v3.0 - December 28, 2024 (FINAL FIXED)
  - CRITICAL FIX: Two-pass parsing for achievement standards
  - Added nested table support (handles Media Arts VCAMAP044 issue)

CSV OUTPUT STRUCTURE (ALIGNED WITH ACARA v4.2):
-----------------------------------------------
Curriculum Files: SubjectArea, Subject, Band, Strand, Substrand, ContentCode, 
                  ContentDescription, Topic, Elaboration

Achievement Standard Components: SubjectArea, Subject, Band, ASComponentCode, ASComponentText,
                                 Strand, Keywords, LinkedContentCodes, ConfidenceScore

Glossary: SubjectArea, Subject, Band, Term, Definition

BAND NAME TRANSFORMATION (v4.5):
--------------------------------
ALL VCAA documents use "Level" internally. CSV outputs convert to "Year" for
Australian school terminology - EXCEPT EAL which keeps "Level" for pathway levels.

Document Heading          -> CSV Output
'Foundation Level A'      -> 'Foundation A'      (Maths pathways)
'Foundation Level B'      -> 'Foundation B'
'Foundation Level C'      -> 'Foundation C'
'Foundation Level D'      -> 'Foundation D'
'Foundation'              -> 'Foundation'
'Level 1'                 -> 'Year 1'
'Level 2'                 -> 'Year 2'
... (through to)
'Level 10'                -> 'Year 10'
'Level 10A'               -> 'Year 10A'
'Levels 1 and 2'          -> 'Years 1 and 2'     (Arts subjects)
'Levels 3 and 4'          -> 'Years 3 and 4'
'Levels 5 and 6'          -> 'Years 5 and 6'
'Levels 7 and 8'          -> 'Years 7 and 8'
'Levels 9 and 10'         -> 'Years 9 and 10'
'Level AL'                -> 'Level AL'          (EAL ONLY - exception)
'Level A1'                -> 'Level A1'          (EAL ONLY)
'Level BL'                -> 'Level BL'          (EAL ONLY)
'Level C4'                -> 'Level C4'          (EAL ONLY)

SUPPORTED BAND CONFIGURATIONS:
------------------------------
Standard subjects: Foundation, Year 1-10
Mathematics: Foundation A/B/C/D, Foundation, Year 1-10, Year 10A
Arts subjects: Foundation, Years 1 and 2, Years 3 and 4, Years 5 and 6, Years 7 and 8, Years 9 and 10
EAL (English as an Additional Language): Level AL, A1, A2, BL, B1, B2, B3, CL, C1, C2, C3, C4
"""

import re
import csv
import os
import sys
from dataclasses import dataclass, field
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from typing import List, Dict, Set, Optional, Tuple
from collections import defaultdict

# Ensure emoji/Unicode prints don't crash Windows cp1252 consoles
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    pass


# ============================================================================
# VERSION INFO
# ============================================================================

PARSER_VERSION = "4.5.0"
PARSER_DATE = "2026-02-01"


# ============================================================================
# CONFIGURATION - CHANGE THESE FOR YOUR SUBJECT
# ============================================================================

# EXAMPLE CONFIGURATIONS:
#
# Standard subject (English, Science, History, etc.):
#   'bands': ["Foundation", "Level 1", "Level 2", ..., "Level 10"]
#
# Mathematics (with Foundation pathways and 10A):
#   'bands': ["Foundation Level A", "Foundation Level B", "Foundation Level C", 
#             "Foundation Level D", "Foundation", "Level 1", ..., "Level 10", "Level 10A"]
#
# Arts subjects (Dance, Drama, Music, Visual Arts, etc.):
#   'bands': ["Foundation", "Levels 1 and 2", "Levels 3 and 4", "Levels 5 and 6",
#             "Levels 7 and 8", "Levels 9 and 10"]
#
# EAL (English as an Additional Language) - uses pathway levels:
#   'bands': ["Level AL", "Level A1", "Level A2", "Level BL", "Level B1", "Level B2", 
#             "Level B3", "Level CL", "Level C1", "Level C2", "Level C3", "Level C4"]

SUBJECT_CONFIG = {
    'subject_area': 'English',  # e.g., 'English', 'Mathematics', 'The Arts', 'Science'
    'subject': 'English',       # e.g., 'English', 'Mathematics', 'Dance', 'Science'
    'bands': [
        "Foundation",
        "Level 1", "Level 2", "Level 3", "Level 4", "Level 5",
        "Level 6", "Level 7", "Level 8", "Level 9", "Level 10"
    ],
    'strands': None,  # Will be discovered from document if None
    'curriculum_doc': '/path/to/curriculum.docx',
    'comparison_doc': None,  # Optional: '/path/to/comparison.docx'
    'output_dir': '/path/to/output'
}

# Pattern must handle:
# - Standard: VC2M7N01 (Maths Level 7 Number 01)
# - Foundation pathways: VC2MFAN01, VC2MFBN01, VC2MFCN01, VC2MFDN01 
# - Arts: VC2ADAFE01, VC2ADA6C01, VC2ADA10P03
# - Two-digit levels: VC2M10A01
# - EAL codes: VC2EALA1L01, VC2EALALL01, VC2EALCLR10 (Pathway A/B/C, Level L/1/2/3/4, Strand L/R/W)
V2_CODE_PATTERN = re.compile(r'VC2(?:EAL[ABC][L1234][LRW]\d+|[A-Z]+(?:F[A-D]?|\d{1,2})[A-Z]+\d+)')


# ============================================================================
# DATA CLASSES (ALIGNED WITH ACARA)
# ============================================================================

@dataclass
class ContentDescriptor:
    """Represents a single curriculum content descriptor."""
    code: str
    description: str
    subject_area: str
    subject: str
    band: str
    strand: str
    substrand: str
    topic: str = ""
    elaborations: List[str] = field(default_factory=list)


@dataclass
class AchievementStandardComponent:
    """Represents a parsed component of an achievement standard."""
    code: str  # ALIGNED: Added ASComponentCode for linking (e.g., VC2EFASC01)
    subject_area: str
    subject: str
    band: str
    text: str
    strand: str
    keywords: str
    linked_codes: str
    confidence: str


@dataclass
class AchievementStandardComparison:
    """Represents V1-V2 achievement standard comparison."""
    band: str
    achievement_v1: str
    achievement_v2: str
    comment: str


@dataclass
class CurriculumComparison:
    """Represents V1-V2 curriculum content comparison."""
    band: str
    code_v2: str
    code_v1: str
    description_v1: str
    change_type: str
    comment: str


# ============================================================================
# STANDARDISED KEYWORD EXTRACTION (ALIGNED WITH ACARA)
# ============================================================================

STOP_WORDS = {
    'the', 'and', 'of', 'to', 'a', 'an', 'in', 'on', 'at', 'for', 'with',
    'as', 'by', 'from', 'or', 'is', 'are', 'be', 'been', 'being',
    'that', 'this', 'these', 'those', 'they', 'them', 'their', 'it', 'its',
    'has', 'have', 'had', 'was', 'were', 'will', 'would', 'could', 'should',
    'may', 'might', 'must', 'shall', 'can', 'do', 'does', 'did',
    'how', 'which', 'also', 'than', 'then', 'end', 'both', 'each', 'such',
    'into', 'through', 'students', 'including', 'using'
}


def extract_keywords(text: str) -> Set[str]:
    """Extract meaningful keywords from text. Min word length 3, no hard limit."""
    if not text:
        return set()
    words = text.lower().replace('.', '').replace(',', '').replace(';', '').split()
    return {w for w in words if len(w) > 3 and w not in STOP_WORDS}


# ============================================================================
# GENERIC HELPER FUNCTIONS
# ============================================================================

# US to AU English spelling conversions
# Covers common -ize/-ise, -or/-our, -er/-re, -og/-ogue patterns
US_TO_AU_SPELLING = {
    # -ize to -ise (and derivatives)
    'recognize': 'recognise', 'recognizes': 'recognises', 'recognized': 'recognised', 'recognizing': 'recognising',
    'organize': 'organise', 'organizes': 'organises', 'organized': 'organised', 'organizing': 'organising',
    'analyze': 'analyse', 'analyzes': 'analyses', 'analyzed': 'analysed', 'analyzing': 'analysing',
    'summarize': 'summarise', 'summarizes': 'summarises', 'summarized': 'summarised', 'summarizing': 'summarising',
    'categorize': 'categorise', 'categorizes': 'categorises', 'categorized': 'categorised', 'categorizing': 'categorising',
    'prioritize': 'prioritise', 'prioritizes': 'prioritises', 'prioritized': 'prioritised', 'prioritizing': 'prioritising',
    'utilize': 'utilise', 'utilizes': 'utilises', 'utilized': 'utilised', 'utilizing': 'utilising',
    'minimize': 'minimise', 'minimizes': 'minimises', 'minimized': 'minimised', 'minimizing': 'minimising',
    'maximize': 'maximise', 'maximizes': 'maximises', 'maximized': 'maximised', 'maximizing': 'maximising',
    'emphasize': 'emphasise', 'emphasizes': 'emphasises', 'emphasized': 'emphasised', 'emphasizing': 'emphasising',
    'visualize': 'visualise', 'visualizes': 'visualises', 'visualized': 'visualised', 'visualizing': 'visualising',
    'realize': 'realise', 'realizes': 'realises', 'realized': 'realised', 'realizing': 'realising',
    'normalize': 'normalise', 'normalizes': 'normalises', 'normalized': 'normalised', 'normalizing': 'normalising',
    'standardize': 'standardise', 'standardizes': 'standardises', 'standardized': 'standardised', 'standardizing': 'standardising',
    'customize': 'customise', 'customizes': 'customises', 'customized': 'customised', 'customizing': 'customising',
    'specialize': 'specialise', 'specializes': 'specialises', 'specialized': 'specialised', 'specializing': 'specialising',
    'generalize': 'generalise', 'generalizes': 'generalises', 'generalized': 'generalised', 'generalizing': 'generalising',
    'memorize': 'memorise', 'memorizes': 'memorises', 'memorized': 'memorised', 'memorizing': 'memorising',
    'theorize': 'theorise', 'theorizes': 'theorises', 'theorized': 'theorised', 'theorizing': 'theorising',
    'hypothesize': 'hypothesise', 'hypothesizes': 'hypothesises', 'hypothesized': 'hypothesised', 'hypothesizing': 'hypothesising',
    'synthesize': 'synthesise', 'synthesizes': 'synthesises', 'synthesized': 'synthesised', 'synthesizing': 'synthesising',
    'criticize': 'criticise', 'criticizes': 'criticises', 'criticized': 'criticised', 'criticizing': 'criticising',
    'apologize': 'apologise', 'apologizes': 'apologises', 'apologized': 'apologised', 'apologizing': 'apologising',
    # -or to -our
    'color': 'colour', 'colors': 'colours', 'colored': 'coloured', 'coloring': 'colouring',
    'behavior': 'behaviour', 'behaviors': 'behaviours',
    'favor': 'favour', 'favors': 'favours', 'favored': 'favoured', 'favoring': 'favouring', 'favorable': 'favourable',
    'honor': 'honour', 'honors': 'honours', 'honored': 'honoured', 'honoring': 'honouring',
    'labor': 'labour', 'labors': 'labours', 'labored': 'laboured', 'laboring': 'labouring',
    'neighbor': 'neighbour', 'neighbors': 'neighbours', 'neighboring': 'neighbouring',
    'humor': 'humour', 'humors': 'humours',
    'vigor': 'vigour',
    'endeavor': 'endeavour', 'endeavors': 'endeavours',
    # -er to -re
    'center': 'centre', 'centers': 'centres', 'centered': 'centred', 'centering': 'centring',
    'meter': 'metre', 'meters': 'metres',
    'theater': 'theatre', 'theaters': 'theatres',
    'fiber': 'fibre', 'fibers': 'fibres',
    # -og to -ogue
    'catalog': 'catalogue', 'catalogs': 'catalogues',
    'dialog': 'dialogue', 'dialogs': 'dialogues',
    'analog': 'analogue',
    'prolog': 'prologue',
    'epilog': 'epilogue',
    # Other common differences
    'program': 'programme', 'programs': 'programmes',  # Note: 'program' is OK for computer programs
    'defense': 'defence',
    'offense': 'offence',
    'license': 'licence',  # noun form
    'practice': 'practise',  # verb form - be careful with this one
    'judgment': 'judgement',
    'aging': 'ageing',
    'modeling': 'modelling',
    'traveling': 'travelling', 'traveled': 'travelled', 'traveler': 'traveller',
    'canceled': 'cancelled', 'canceling': 'cancelling',
    'labeled': 'labelled', 'labeling': 'labelling',
    'leveled': 'levelled', 'leveling': 'levelling',
    'signaled': 'signalled', 'signaling': 'signalling',
}


def convert_to_au_english(text: str) -> str:
    """Convert US English spellings to Australian English.
    
    Processes text to ensure consistent AU English spelling in output.
    Case-insensitive matching with case preservation.
    """
    if not text:
        return text
    
    result = text
    for us_spelling, au_spelling in US_TO_AU_SPELLING.items():
        # Case-insensitive replacement with case preservation
        pattern = re.compile(re.escape(us_spelling), re.IGNORECASE)
        
        def replace_preserve_case(match):
            matched = match.group(0)
            if matched.isupper():
                return au_spelling.upper()
            elif matched[0].isupper():
                return au_spelling.capitalize()
            return au_spelling
        
        result = pattern.sub(replace_preserve_case, result)
    
    return result


def capitalise_first(text: str) -> str:
    """Capitalise first letter of text while preserving the rest."""
    if not text:
        return ""
    return text[0].upper() + text[1:] if len(text) > 1 else text.upper()


def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace while preserving line breaks and converting to AU English."""
    if not text:
        return ""
    text = re.sub(r'[^\S\n]+', ' ', text)
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    cleaned = '\n'.join(lines)
    # Apply AU English spelling conversion
    return convert_to_au_english(cleaned)


def get_cell_text(cell) -> str:
    """Extract text from a table cell, handling nested tables."""
    cell_text = cell.text.strip()
    if not cell_text:
        from docx.table import Table
        nested_tables = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
        if nested_tables:
            texts = []
            for tbl_element in nested_tables:
                nested_table = Table(tbl_element, cell._parent._parent._parent)
                for row in nested_table.rows:
                    for nested_cell in row.cells:
                        text = nested_cell.text.strip()
                        if text:
                            texts.append(text)
            cell_text = '\n'.join(texts)
    return cell_text


def get_full_paragraph_text(para) -> str:
    """Extract all text from paragraph including from nested XML elements.
    
    Some VCAA documents have text in nested w:t elements that para.text misses.
    This extracts text from all w:t elements in the paragraph's XML.
    """
    from docx.oxml.ns import qn
    text_elements = para._element.findall('.//' + qn('w:t'))
    if text_elements:
        return ''.join([t.text for t in text_elements if t.text])
    return para.text or ''


def extract_v1_codes(text: str) -> List[str]:
    """Extract V1 curriculum codes from text."""
    if not text:
        return []
    codes = re.findall(r'\(VC[A-Z]{2,4}\d+\)?', text)
    return [c.strip('()') for c in codes]


def extract_v2_code(text: str) -> str:
    """Extract V2 curriculum code from text."""
    if not text:
        return ""
    match = V2_CODE_PATTERN.search(text)
    return match.group(0) if match else ""


def normalise_band_name(text: str, valid_bands: List[str]) -> str:
    """Normalise band name to handle Level vs Levels."""
    if text in valid_bands:
        return text
    if text.startswith("Level ") and not text.startswith("Levels "):
        normalised = "Levels " + text[6:]
        if normalised in valid_bands:
            return normalised
    if text.startswith("Levels "):
        normalised = "Level " + text[7:]
        if normalised in valid_bands:
            for band in valid_bands:
                if band.endswith(text[7:]):
                    return band
    return text


def transform_band_for_output(band: str) -> str:
    """Transform band name from VCAA document format to CSV/database format.
    
    IMPORTANT: VCAA documents ALL use "Level" terminology internally.
    CSV outputs use "Year" terminology to match Australian school conventions.
    
    =========================================================================
    DOCUMENT FORMAT (input)          ->  CSV FORMAT (output)
    =========================================================================
    
    STANDARD SUBJECTS (English, Maths, Science, History, etc.):
      'Foundation'                   ->  'Foundation'
      'Foundation Level A'           ->  'Foundation A'      (Maths pathways)
      'Foundation Level B'           ->  'Foundation B'
      'Foundation Level C'           ->  'Foundation C'
      'Foundation Level D'           ->  'Foundation D'
      'Level 1'                      ->  'Year 1'
      'Level 2'                      ->  'Year 2'
      ... (through to Level 10)
      'Level 10'                     ->  'Year 10'
      'Level 10A'                    ->  'Year 10A'          (Maths extension)
    
    ARTS SUBJECTS (Dance, Drama, Music, Visual Arts, Media Arts, VCD):
      'Levels 1 and 2'               ->  'Years 1 and 2'
      'Levels 3 and 4'               ->  'Years 3 and 4'
      'Levels 5 and 6'               ->  'Years 5 and 6'
      'Levels 7 and 8'               ->  'Years 7 and 8'
      'Levels 9 and 10'              ->  'Years 9 and 10'
    
    EAL - EXCEPTION (English as an Additional Language):
      EAL is the ONLY subject that keeps "Level" in the CSV output.
      This is because EAL uses proficiency pathway levels, not year levels.
      A student can be in Year 7 but at EAL Level B1 based on their 
      English proficiency. EAL runs alongside regular year-level classes.
      
      'Level AL'                     ->  'Level AL'   (Pathway A - Literacy)
      'Level A1'                     ->  'Level A1'   (Pathway A - Level 1)
      'Level A2'                     ->  'Level A2'   (Pathway A - Level 2)
      'Level BL'                     ->  'Level BL'   (Pathway B - Literacy)
      'Level B1'                     ->  'Level B1'   (Pathway B - Level 1)
      'Level B2'                     ->  'Level B2'   (Pathway B - Level 2)
      'Level B3'                     ->  'Level B3'   (Pathway B - Level 3)
      'Level CL'                     ->  'Level CL'   (Pathway C - Literacy)
      'Level C1'                     ->  'Level C1'   (Pathway C - Level 1)
      'Level C2'                     ->  'Level C2'   (Pathway C - Level 2)
      'Level C3'                     ->  'Level C3'   (Pathway C - Level 3)
      'Level C4'                     ->  'Level C4'   (Pathway C - Level 4)
    =========================================================================
    """
    if not band:
        return band
    
    # -------------------------------------------------------------------------
    # EAL EXCEPTION: Keep "Level" for EAL pathway levels (the ONLY exception)
    # Pattern: Level + pathway letter (A/B/C) + level indicator (L/1/2/3/4)
    # -------------------------------------------------------------------------
    eal_pattern = re.compile(r'^Level [ABC][L1234]$')
    if eal_pattern.match(band):
        return band  # Keep as "Level AL", "Level B3", etc.
    
    # -------------------------------------------------------------------------
    # ALL OTHER SUBJECTS: Convert "Level(s)" to "Year(s)"
    # -------------------------------------------------------------------------
    
    # Maths Foundation pathways: "Foundation Level X" -> "Foundation X"
    if band.startswith("Foundation Level "):
        return "Foundation " + band[17:]
    
    # Arts subjects: "Levels X and Y" -> "Years X and Y"
    if band.startswith("Levels "):
        return "Years " + band[7:]  # Replace "Levels " with "Years "
    
    # Standard subjects: "Level X" -> "Year X"
    if band.startswith("Level "):
        return "Year " + band[6:]  # Replace "Level " with "Year "
    
    # Foundation and any other formats - unchanged
    return band


# ============================================================================
# PRE-CHECK TESTS
# ============================================================================

class VictorianPrecheckTests:
    """Pre-parsing validation tests."""
    
    def __init__(self, curriculum_doc: str, comparison_doc: str = None):
        self.curriculum_doc = curriculum_doc
        self.comparison_doc = comparison_doc
        self.issues = []
        self.warnings = []
        self.passed_tests = 0
        self.total_tests = 0
    
    def run_all(self) -> bool:
        print("\n" + "="*80)
        print("PHASE 1: PRE-CHECK VALIDATION")
        print("="*80)
        
        self._check_document_access()
        self._check_document_structure()
        self._check_for_nested_tables()
        self._check_code_patterns()
        
        print("\n" + "-"*80)
        print(f"Tests passed: {self.passed_tests}/{self.total_tests}")
        if self.issues:
            print(f"\n❌ CRITICAL ISSUES: {self.issues}")
        if self.warnings:
            print(f"\n⚠️  WARNINGS: {self.warnings}")
        print("="*80)
        
        return len(self.issues) == 0
    
    def _check_document_access(self):
        self.total_tests += 1
        print("\n1. DOCUMENT ACCESS")
        docs = [self.curriculum_doc]
        if self.comparison_doc:
            docs.append(self.comparison_doc)
        accessible = 0
        for doc_path in docs:
            try:
                Document(doc_path)
                print(f"   ✅ {Path(doc_path).name}")
                accessible += 1
            except Exception as e:
                self.issues.append(f"Cannot access {doc_path}: {e}")
                print(f"   ❌ {doc_path}: ERROR")
        if accessible == len(docs):
            self.passed_tests += 1
    
    def _check_document_structure(self):
        self.total_tests += 1
        print("\n2. DOCUMENT STRUCTURE")
        try:
            doc = Document(self.curriculum_doc)
            print(f"   Tables: {len(doc.tables)}, Paragraphs: {len(doc.paragraphs)}")
            if len(doc.tables) > 0:
                self.passed_tests += 1
            else:
                self.issues.append("No tables found")
        except Exception as e:
            self.issues.append(f"Structure check failed: {e}")
    
    def _check_for_nested_tables(self):
        self.total_tests += 1
        print("\n3. NESTED TABLES SCAN")
        try:
            doc = Document(self.curriculum_doc)
            nested_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        nested = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
                        if nested:
                            nested_count += len(nested)
            if nested_count > 0:
                print(f"   ℹ️  {nested_count} nested tables found (get_cell_text will handle)")
            else:
                print("   ✅ No nested tables")
            self.passed_tests += 1
        except Exception as e:
            self.warnings.append(f"Nested table scan failed: {e}")
    
    def _check_code_patterns(self):
        self.total_tests += 1
        print("\n4. V2 CODE PATTERN CHECK")
        try:
            doc = Document(self.curriculum_doc)
            codes_found = set()
            for table in doc.tables[:10]:
                for row in table.rows:
                    for cell in row.cells:
                        codes_found.update(V2_CODE_PATTERN.findall(cell.text))
            if codes_found:
                print(f"   ✅ Found {len(codes_found)} unique V2 codes")
                self.passed_tests += 1
            else:
                self.warnings.append("No V2 codes found in sample")
                print("   ⚠️  No V2 codes detected")
        except Exception as e:
            self.warnings.append(f"Code pattern check failed: {e}")


# ============================================================================
# POST-PARSE TESTS
# ============================================================================

class VictorianParsedDataTests:
    """Tests run after parsing to validate data quality."""
    
    def __init__(self, content_descriptors: List[ContentDescriptor], 
                 as_components: List[AchievementStandardComponent]):
        self.content_descriptors = content_descriptors
        self.as_components = as_components
        self.test_results = []
    
    def run_all(self) -> bool:
        print("\n" + "="*80)
        print("PHASE 3: POST-PARSE DATA VALIDATION")
        print("="*80)
        
        self._test_code_extraction()
        self._test_capitalisation()
        self._test_strand_inference()
        self._test_elaboration_counts()
        self._test_no_duplicates()
        
        print("\n" + "-"*80)
        passed = sum(1 for r in self.test_results if r['passed'])
        print(f"Total: {passed}/{len(self.test_results)} tests passed")
        print("="*80)
        
        return all(r['passed'] for r in self.test_results)
    
    def _test_code_extraction(self):
        print("\n1. CODE EXTRACTION")
        invalid = [cd.code for cd in self.content_descriptors if not V2_CODE_PATTERN.match(cd.code)]
        passed = len(invalid) == 0
        self.test_results.append({'test': 'Code Extraction', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(self.content_descriptors)} codes, {len(invalid)} invalid")
    
    def _test_capitalisation(self):
        print("\n2. CAPITALISATION")
        uncap_desc = [cd for cd in self.content_descriptors if cd.description and cd.description[0].islower()]
        uncap_elab = sum(1 for cd in self.content_descriptors for e in cd.elaborations if e and e[0].islower())
        passed = len(uncap_desc) == 0 and uncap_elab == 0
        self.test_results.append({'test': 'Capitalisation', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} Descriptions: {len(uncap_desc)} uncapitalised, Elaborations: {uncap_elab}")
    
    def _test_strand_inference(self):
        print("\n3. STRAND INFERENCE")
        no_strand = [c for c in self.as_components if not c.strand]
        passed = len(no_strand) == 0
        self.test_results.append({'test': 'Strand Inference', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(self.as_components)} components, {len(no_strand)} missing strand")
    
    def _test_elaboration_counts(self):
        print("\n4. ELABORATIONS")
        total = sum(len(cd.elaborations) for cd in self.content_descriptors)
        with_elab = sum(1 for cd in self.content_descriptors if cd.elaborations)
        passed = total > 0
        self.test_results.append({'test': 'Elaborations', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {total} total, {with_elab} CDs have elaborations")
    
    def _test_no_duplicates(self):
        print("\n5. DUPLICATES")
        codes = [cd.code for cd in self.content_descriptors]
        dups = [c for c in set(codes) if codes.count(c) > 1]
        passed = len(dups) == 0
        self.test_results.append({'test': 'Duplicates', 'passed': passed})
        print(f"   {'✅' if passed else '❌'} {len(set(codes))} unique codes, {len(dups)} duplicates")




# ============================================================================
# CURRICULUM FILE PARSER
# ============================================================================

class CurriculumParser:
    """Parses Victorian Curriculum V2.0 documents to extract content descriptors."""
    
    def __init__(self, config: Dict):
        self.config = config
        self.subject_area = config['subject_area']
        self.subject = config['subject']
        self.bands = config['bands']
        self.doc = Document(config['curriculum_doc'])
        self.content_descriptors: List[ContentDescriptor] = []
        self.as_components: List[AchievementStandardComponent] = []
        self.band_component_counts = defaultdict(int)  # ALIGNED: Track component counts per band
        
        if config.get('strands') is None:
            print("\nDISCOVERING STRANDS FROM DOCUMENT")
            self.strands = self._discover_strands()
            print(f"Found {len(self.strands)} strands")
        else:
            self.strands = config['strands']
    
    def _get_band_code(self, band: str) -> str:
        """Convert band name to code for ASComponentCode generation.
        
        ALIGNED: Same logic as ACARA parser.
        Examples:
            'Foundation' -> 'F'
            'Levels 1 and 2' -> 'L1-2'
            'Levels 3 and 4' -> 'L3-4'
            'Level AL' -> 'AL' (EAL pathway)
            'Level B3' -> 'B3' (EAL pathway)
        """
        if band == 'Foundation':
            return 'F'
        
        # Handle EAL pathway levels (Level AL, A1, A2, BL, B1, B2, B3, CL, C1, C2, C3, C4)
        eal_match = re.match(r'Level ([ABC][L1234])$', band)
        if eal_match:
            return eal_match.group(1)
        
        # Handle "Levels X and Y" format
        match = re.match(r'Levels? (\d+)(?: and (\d+))?', band)
        if match:
            start = match.group(1)
            end = match.group(2)
            if end:
                return f"L{start}-{end}"
            return f"L{start}"
        return band[0].upper()  # Fallback: first character
    
    def _generate_asc_code(self, band: str) -> str:
        """Generate ASComponentCode for Victorian curriculum.
        
        Format: VC2{subject_prefix}{band_code}ASC{sequence:02d}
        
        For EAL: VC2EAL{level}ASC{seq} e.g., VC2EALALASC01, VC2EALA1ASC01, VC2EALB3ASC01
        For other subjects: VC2{prefix}{band_code}ASC{seq} e.g., VC2EL1-2ASC01
        """
        # Get subject prefix from first content descriptor code, or use first 2 letters of subject
        if self.content_descriptors:
            sample_code = self.content_descriptors[0].code
            # Extract subject letters from code like VC2EART12P01 -> EART or VC2EALA1L01 -> EAL
            match = re.match(r'VC2([A-Z]+)', sample_code)
            subject_prefix = match.group(1)[:3] if match else self.subject[:2].upper()
        else:
            subject_prefix = self.subject[:2].upper()
        
        # Handle EAL - extract level from band (e.g., "Level AL" -> "AL", "Level B3" -> "B3")
        eal_match = re.match(r'Level ([ABC][L1234])$', band)
        if eal_match:
            level_code = eal_match.group(1)  # AL, A1, A2, BL, B1, B2, B3, CL, C1, C2, C3, C4
            self.band_component_counts[band] += 1
            return f"VC2EAL{level_code}ASC{self.band_component_counts[band]:02d}"
        
        # For other subjects, use existing logic
        band_code = self._get_band_code(band)
        self.band_component_counts[band] += 1
        return f"VC2{subject_prefix}{band_code}ASC{self.band_component_counts[band]:02d}"
    
    def _discover_strands(self) -> Dict[str, List[str]]:
        discovered = {}
        current_strand = None
        for para in self.doc.paragraphs:
            text = clean_text(para.text)
            if text.startswith("Strand: "):
                current_strand = text.replace("Strand: ", "").strip()
                if current_strand not in discovered:
                    discovered[current_strand] = []
                    print(f"  Strand: {current_strand}")
            elif text.startswith("Sub-strand: ") and current_strand:
                substrand = text.replace("Sub-strand: ", "").strip()
                if substrand not in discovered[current_strand]:
                    discovered[current_strand].append(substrand)
        return discovered
    
    def parse_curriculum(self):
        """Two-pass parsing for accurate strand inference."""
        current_band = None
        current_strand = None
        current_substrand = None
        achievement_paragraphs_by_band = {}
        collecting_achievement = False
        current_achievement_paragraphs = []  # List of (text, strand) tuples
        current_as_strand = None  # Track strand within AS section
        tables_by_element = {id(table._element): table for table in self.doc.tables}
        
        # Known strand names for AS section headings
        as_strand_names = {'Listening and Speaking', 'Reading and Viewing', 'Writing'}
        
        print("\n=== PASS 1: Parsing Curriculum Content ===")
        
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                for para in self.doc.paragraphs:
                    if para._element == element:
                        # Use get_full_paragraph_text to handle nested XML text elements
                        # Some VCAA docs have text that para.text misses
                        text = clean_text(get_full_paragraph_text(para))
                        if not text:
                            continue
                        
                        band = normalise_band_name(text, self.bands)
                        if band in self.bands:
                            if collecting_achievement and current_achievement_paragraphs and current_band:
                                achievement_paragraphs_by_band[current_band] = current_achievement_paragraphs
                                current_achievement_paragraphs = []
                            current_band = band
                            collecting_achievement = False
                            current_as_strand = None
                            print(f"\n  Band: {current_band}")
                            break
                        
                        if text == "Achievement standard":
                            collecting_achievement = True
                            current_achievement_paragraphs = []
                            current_as_strand = None
                            break
                        
                        if collecting_achievement:
                            if text.startswith("Content descriptions") or text.startswith("Strand:"):
                                if current_achievement_paragraphs and current_band:
                                    achievement_paragraphs_by_band[current_band] = current_achievement_paragraphs
                                collecting_achievement = False
                                current_achievement_paragraphs = []
                                current_as_strand = None
                            elif text in as_strand_names:
                                # This is a strand heading within the AS section
                                current_as_strand = text
                                break
                            elif text.startswith("By the end of") or text.startswith("Students") or text.startswith("They"):
                                # Store text with its strand
                                current_achievement_paragraphs.append((text, current_as_strand))
                                break
                        
                        strand_check = text.replace("Strand: ", "").strip() if text.startswith("Strand: ") else text
                        if strand_check in self.strands.keys():
                            current_strand = strand_check
                            break
                        
                        substrand_check = text.replace("Sub-strand: ", "").strip() if text.startswith("Sub-strand: ") else text
                        for strand, substrands in self.strands.items():
                            if substrand_check in substrands:
                                current_substrand = substrand_check
                                break
                        break
            
            elif element.tag.endswith('tbl'):
                table = tables_by_element.get(id(element))
                if table and current_band:
                    self._parse_curriculum_table(table, current_band, current_strand, current_substrand)
        
        if collecting_achievement and current_achievement_paragraphs and current_band:
            achievement_paragraphs_by_band[current_band] = current_achievement_paragraphs
        
        print(f"\n=== PASS 2: Parsing Achievement Standards ===")
        print(f"Total content descriptors: {len(self.content_descriptors)}")
        
        for band in self.bands:
            if band in achievement_paragraphs_by_band:
                self._parse_achievement_paragraphs(achievement_paragraphs_by_band[band], band)
        
        print(f"Total AS components: {len(self.as_components)}")
        return self.content_descriptors
    
    def _parse_curriculum_table(self, table, band, strand, substrand):
        if len(table.columns) < 2 or len(table.rows) < 2:
            return
        header = clean_text(table.rows[0].cells[0].text).lower()
        if 'content description' not in header:
            return
        
        has_topic = False
        if len(table.rows[0].cells) >= 3:
            if 'topic' in clean_text(table.rows[0].cells[2].text).lower():
                has_topic = True
        
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            if len(row.cells) < 2:
                continue
            
            cd_text = clean_text(row.cells[0].text)
            v2_code = extract_v2_code(cd_text)
            if not v2_code:
                continue
            
            # Remove V2 code from description - use same pattern as V2_CODE_PATTERN
            cd_desc = V2_CODE_PATTERN.sub('', cd_text).strip()
            
            # Handle EAL format: "Students learn to:\ndescription\ncode"
            # Remove the prefix and extract just the description
            if cd_desc.startswith("Students learn to:"):
                cd_desc = cd_desc.replace("Students learn to:", "").strip()
            
            # Clean up the description (remove newlines, get just the content)
            cd_desc = ' '.join(cd_desc.split())  # Normalise whitespace
            cd_desc = capitalise_first(cd_desc)
            
            # Handle elaboration format: "This may involve students:\nelab1\nelab2..."
            elab_paragraphs = [p.text.strip() for p in row.cells[1].paragraphs if p.text.strip()]
            elaborations = []
            for elab in elab_paragraphs:
                # Skip the prefix line
                if elab.lower().startswith("this may involve students:"):
                    continue
                elaborations.append(capitalise_first(elab))
            
            topic = ""
            if has_topic and len(row.cells) >= 3:
                topic = capitalise_first(clean_text(row.cells[2].text))
            
            self.content_descriptors.append(ContentDescriptor(
                code=v2_code,
                description=cd_desc,
                subject_area=self.subject_area,
                subject=self.subject,
                band=band,
                strand=strand or "",
                substrand=substrand or "",
                topic=topic,
                elaborations=elaborations
            ))
            print(f"    {v2_code} ({len(elaborations)} elabs)")
    
    def _parse_achievement_paragraphs(self, paragraphs: List, band: str):
        """Parse achievement standard paragraphs into components.
        
        Args:
            paragraphs: List of (text, strand) tuples where strand is from document structure
            band: The band/level name
        """
        for para_data in paragraphs:
            # Handle both tuple format (text, strand) and legacy string format
            if isinstance(para_data, tuple):
                para, doc_strand = para_data
            else:
                para = para_data
                doc_strand = None
            
            sentences = [s.strip() for s in para.split('. ') if s.strip()]
            for sentence in sentences:
                if not sentence.endswith('.'):
                    sentence += '.'
                keywords = extract_keywords(sentence.lower())
                
                # Use strand from document structure if available, otherwise infer
                if doc_strand:
                    strand = doc_strand
                else:
                    strand = self._infer_strand(sentence)
                
                linked_codes, confidence = self._find_linked_codes(band, strand, keywords)
                
                # ALIGNED: Generate ASComponentCode like ACARA parser
                asc_code = self._generate_asc_code(band)
                
                self.as_components.append(AchievementStandardComponent(
                    code=asc_code,
                    subject_area=self.subject_area,
                    subject=self.subject,
                    band=band,
                    text=capitalise_first(sentence),
                    strand=strand,
                    keywords=','.join(sorted(keywords)),
                    linked_codes=','.join(linked_codes),
                    confidence=confidence
                ))
    
    def _infer_strand(self, text: str) -> str:
        if not self.content_descriptors:
            return list(self.strands.keys())[0] if self.strands else "Unknown"
        
        component_keywords = extract_keywords(text.lower())
        strand_scores = {}
        
        for cd in self.content_descriptors:
            cd_keywords = extract_keywords(cd.description.lower())
            overlap = len(component_keywords & cd_keywords)
            if cd.strand not in strand_scores:
                strand_scores[cd.strand] = 0
            strand_scores[cd.strand] += overlap
        
        if strand_scores and max(strand_scores.values()) > 0:
            return max(strand_scores, key=strand_scores.get)
        return list(self.strands.keys())[0] if self.strands else "Unknown"
    
    def _find_linked_codes(self, band: str, strand: str, keywords: Set[str]) -> Tuple[List[str], str]:
        matched = []
        for cd in self.content_descriptors:
            if cd.band != band:
                continue
            score = 2 if cd.strand == strand else 0
            cd_text = (cd.description + ' ' + ' '.join(cd.elaborations)).lower()
            score += sum(1 for kw in keywords if kw in cd_text)
            if score >= 2:
                matched.append((cd.code, score))
        
        matched.sort(key=lambda x: x[1], reverse=True)
        top_codes = [code for code, _ in matched[:5]]
        
        if matched and matched[0][1] >= 4:
            confidence = "High"
        elif matched and matched[0][1] >= 2:
            confidence = "Medium"
        else:
            confidence = "Low"
        
        return top_codes, confidence
    
    def save_to_csv(self, output_dir: str):
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        by_band = defaultdict(list)
        for cd in self.content_descriptors:
            by_band[cd.band].append(cd)
        
        for band, cds in by_band.items():
            if not cds:
                continue
            # Transform band name for output (e.g., "Level 1" -> "Year 1")
            output_band = transform_band_for_output(band)
            filename = f"Vic {self.subject} V2 - Curriculum - {output_band}.csv"
            filepath = Path(output_dir) / filename
            
            # ALIGNED column order with ACARA
            # Note: Using ContentCode and ContentDescription (not _V2 suffix) for consistency across subjects
            fieldnames = ['SubjectArea', 'Subject', 'Band', 'Strand', 'Substrand', 
                         'ContentCode', 'ContentDescription', 'Topic', 'Elaboration']
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for cd in cds:
                    if cd.elaborations:
                        for elab in cd.elaborations:
                            writer.writerow({
                                'SubjectArea': cd.subject_area, 'Subject': cd.subject,
                                'Band': output_band, 'Strand': cd.strand, 'Substrand': cd.substrand,
                                'ContentCode': cd.code, 'ContentDescription': cd.description,
                                'Topic': cd.topic, 'Elaboration': elab
                            })
                    else:
                        writer.writerow({
                            'SubjectArea': cd.subject_area, 'Subject': cd.subject,
                            'Band': output_band, 'Strand': cd.strand, 'Substrand': cd.substrand,
                            'ContentCode': cd.code, 'ContentDescription': cd.description,
                            'Topic': cd.topic, 'Elaboration': ''
                        })
            
            rows = sum(len(cd.elaborations) if cd.elaborations else 1 for cd in cds)
            print(f"Saved {rows} rows to {filename}")
        
        if self.as_components:
            filename = f"Vic {self.subject} V2 - Achievement Standard Components.csv"
            filepath = Path(output_dir) / filename
            
            # ALIGNED: Column order matches ACARA parser and import patterns skill
            # Using ASComponentText (not Text) for clarity
            fieldnames = ['SubjectArea', 'Subject', 'Band', 'ASComponentCode', 'ASComponentText',
                         'Strand', 'Keywords', 'LinkedContentCodes', 'ConfidenceScore']
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for comp in self.as_components:
                    # Transform band name for output (e.g., "Level 1" -> "Year 1")
                    output_band = transform_band_for_output(comp.band)
                    writer.writerow({
                        'SubjectArea': comp.subject_area, 'Subject': comp.subject,
                        'Band': output_band, 'ASComponentCode': comp.code, 'ASComponentText': comp.text,
                        'Strand': comp.strand, 'Keywords': comp.keywords,
                        'LinkedContentCodes': comp.linked_codes, 'ConfidenceScore': comp.confidence
                    })
            print(f"Saved {len(self.as_components)} AS components to {filename}")




# ============================================================================
# COMPARISON FILE PARSER
# ============================================================================

class ComparisonParser:
    """Parses VCAA comparison documents for V1-V2 mappings."""
    
    def __init__(self, config: Dict):
        self.config = config
        self.subject_area = config['subject_area']
        self.subject = config['subject']
        self.bands = config['bands']
        self.doc = Document(config['comparison_doc'])
        self.mappings: List[CurriculumComparison] = []
        self.achievement_standards: List[AchievementStandardComparison] = []
    
    def parse_comparison_tables(self):
        current_band = None
        para_idx = 0
        table_idx = 0
        
        for element in self.doc.element.body:
            tag = element.tag.split('}')[-1]
            
            if tag == 'p':
                text = self.doc.paragraphs[para_idx].text.strip()
                normalised = normalise_band_name(text, self.bands)
                if normalised in self.bands:
                    current_band = normalised
                    print(f"\n  Processing: {current_band}")
                para_idx += 1
            
            elif tag == 'tbl':
                table = self.doc.tables[table_idx]
                if len(table.columns) >= 3 and current_band:
                    header = table.rows[0].cells[0].text.lower()
                    if 'version 1' in header or 'v1' in header or 'curriculum f' in header:
                        self._parse_comparison_table(table, current_band)
                table_idx += 1
        
        return self.mappings
    
    def _parse_comparison_table(self, table, band):
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            if len(row.cells) < 3:
                continue
            
            v1_text = clean_text(get_cell_text(row.cells[0]))
            v2_text = clean_text(get_cell_text(row.cells[1]))
            comment = clean_text(get_cell_text(row.cells[2]))
            
            v2_code = extract_v2_code(v2_text)
            v1_lines = [line.strip() for line in v1_text.split('\n') if line.strip()]
            
            v1_entries = []
            for line in v1_lines:
                codes = extract_v1_codes(line)
                for code in codes:
                    desc = re.sub(r'\(' + re.escape(code) + r'\)', '', line).strip()
                    v1_entries.append({'code': code, 'description': capitalise_first(desc)})
            
            if not v1_entries and not v2_code:
                continue
            
            if not v1_entries:
                self.mappings.append(CurriculumComparison(band, v2_code, '', '', 'New', comment))
                print(f"    (New) -> {v2_code}")
                continue
            
            if not v2_code:
                for entry in v1_entries:
                    self.mappings.append(CurriculumComparison(
                        band, '', entry['code'], entry['description'], 'Removed', comment))
                    print(f"    {entry['code']} -> (Removed)")
                continue
            
            change_type = "Refined"
            if 'refined' in comment.lower():
                change_type = "Refined"
            elif 'updated' in comment.lower():
                change_type = "Updated"
            
            for entry in v1_entries:
                self.mappings.append(CurriculumComparison(
                    band, v2_code, entry['code'], entry['description'], change_type, comment))
                print(f"    {entry['code']} -> {v2_code} ({change_type})")
    
    def parse_achievement_standards(self):
        print("\n  Parsing Achievement Standards...")
        current_band = None
        para_idx = 0
        table_idx = 0
        
        for element in self.doc.element.body:
            tag = element.tag.split('}')[-1]
            
            if tag == 'p':
                text = self.doc.paragraphs[para_idx].text.strip()
                normalised = normalise_band_name(text, self.bands)
                if normalised in self.bands:
                    current_band = normalised
                para_idx += 1
            
            elif tag == 'tbl':
                table = self.doc.tables[table_idx]
                if len(table.columns) >= 3 and current_band and len(table.rows) > 1:
                    header = table.rows[0].cells[0].text.lower()
                    row1_text = table.rows[1].cells[0].text.lower() if len(table.rows) > 1 else ""
                    
                    if 'achievement standard' in header or 'by the end' in row1_text:
                        row = table.rows[1]
                        av1 = clean_text(get_cell_text(row.cells[0]))
                        av2 = clean_text(get_cell_text(row.cells[1])) if len(row.cells) > 1 else ""
                        comment = clean_text(get_cell_text(row.cells[2])) if len(row.cells) > 2 else ""
                        
                        if av1.startswith("By the end") or av2.startswith("By the end"):
                            self.achievement_standards.append(
                                AchievementStandardComparison(current_band, av1, av2, comment))
                            print(f"    AS: {current_band}")
                table_idx += 1
        
        print(f"  Total: {len(self.achievement_standards)}")
        return self.achievement_standards
    
    def save_to_csv(self, output_dir: str):
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Curriculum comparison
        filename = f"Vic {self.subject} V2 - Curriculum Comparison - Levels 3 to Year 10.csv"
        filepath = Path(output_dir) / filename
        fieldnames = ['Band', 'ContentCode_V2', 'ContentCode_V1', 'ContentDescription_V1', 'ChangeType', 'Comment']
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for m in self.mappings:
                writer.writerow({
                    'Band': m.band, 'ContentCode_V2': m.code_v2, 'ContentCode_V1': m.code_v1,
                    'ContentDescription_V1': m.description_v1, 'ChangeType': m.change_type, 'Comment': m.comment
                })
        print(f"\nSaved {len(self.mappings)} comparison rows to {filename}")
        
        # Achievement standards
        if self.achievement_standards:
            as_filename = f"Vic {self.subject} V2 - Achievement Standards and Comparison - Foundation to Year 10.csv"
            as_filepath = Path(output_dir) / as_filename
            as_fields = ['Band', 'AchievementStandard_V1', 'AchievementStandard_V2', 'Comment']
            
            with open(as_filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=as_fields)
                writer.writeheader()
                for ast in self.achievement_standards:
                    writer.writerow({
                        'Band': ast.band, 'AchievementStandard_V1': ast.achievement_v1,
                        'AchievementStandard_V2': ast.achievement_v2, 'Comment': ast.comment
                    })
            print(f"Saved {len(self.achievement_standards)} achievement standards to {as_filename}")


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def main():
    config = SUBJECT_CONFIG
    
    print("="*80)
    print(f"Victorian Curriculum V2.0 Parser v{PARSER_VERSION}")
    print(f"Subject: {config['subject']}")
    print("="*80)
    
    # Phase 1: Pre-check
    precheck = VictorianPrecheckTests(config['curriculum_doc'], config.get('comparison_doc'))
    if not precheck.run_all():
        print("\n❌ CANCELLED - Fix issues first")
        return
    
    print("\n✅ Pre-checks passed\n")
    
    # Phase 2: Parse comparison
    print("\n" + "="*60)
    print("PARSING CURRICULUM COMPARISON")
    print("="*60)
    try:
        comp_parser = ComparisonParser(config)
        comp_parser.parse_comparison_tables()
        comp_parser.parse_achievement_standards()
        comp_parser.save_to_csv(config['output_dir'])
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
    
    # Phase 2: Parse curriculum
    print("\n" + "="*60)
    print("PARSING CURRICULUM CONTENT")
    print("="*60)
    curriculum_parser = None
    try:
        curriculum_parser = CurriculumParser(config)
        curriculum_parser.parse_curriculum()
        curriculum_parser.save_to_csv(config['output_dir'])
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
    
    # Phase 3: Post-parse validation
    if curriculum_parser and curriculum_parser.content_descriptors:
        postcheck = VictorianParsedDataTests(
            curriculum_parser.content_descriptors, curriculum_parser.as_components)
        postcheck.run_all()
    
    print("\n" + "="*80)
    print("COMPLETE!")
    print(f"Output: {config['output_dir']}/")
    print("="*80)


if __name__ == "__main__":
    main()
