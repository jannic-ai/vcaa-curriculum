"""
VCE Exam Report Parser — config-driven, subject-agnostic.

Reads a pipeline config JSON from sys.argv[1], locates the exam report DOCX
under VCE_ROOT/<parent_folder>/Documentation/, and emits two CSVs into
ETL_OUTPUT_DIR (or VCE_ROOT/<parent_folder>/CSV):

  vcaa_vce_sd_{slug}_exam_report_overview.csv
  vcaa_vce_sd_{slug}_exam_report_strategies.csv

Config contract:
  documents.exam_reports = { "YYYY": "YYYY-vce-{subject}-report.docx", ... }
  exam_sections = {
      "A": {
          "name": "...",
          "exam_report_code": "VCE{code}ERA",
          "unit_as_code": "...,...",
          "section_type": "multiple_choice" | "written",  # default "written"
          "strategy_format": "prose",
      },
      ...
  }

Algorithm:
1. Find exam sections by matching Heading-2 paragraphs "Section X ..." against
   keys in exam_sections. If no matches and exam_sections has exactly one
   entry, treat the whole document as that section.
2. Everything before the first section heading is "report overview"
   (General Comments / Specific Information). Emit overview rows keyed by
   sub-heading.
3. Within each section:
   - Multiple-choice sections: find the MCQ results table, emit one
     ExaminerCommentary strategy row per question using the Comments cell.
   - Written sections: split into questions by Heading-3 "Question N".
     Classify each paragraph into ExaminerCommentary / EnhancedStrategies /
     LimitedStrategies / StudentResponse using content patterns, and emit a
     strategy row per logical chunk.
"""

import sys
import os
import re
import csv
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Iterable

# Windows cp1252 consoles choke on Unicode prints — force UTF-8 before anything
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    pass

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# ============================================================================
# CONFIG LOADING
# ============================================================================

def load_config() -> Dict:
    if len(sys.argv) < 2:
        print("ERROR: usage: exam-report-parser-template.py <config.json>", file=sys.stderr)
        sys.exit(2)
    config_path = sys.argv[1]
    if not Path(config_path).is_file():
        print(f"ERROR: Config file not found: {config_path}", file=sys.stderr)
        sys.exit(2)
    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)

    vce_root_env = os.environ.get('VCE_ROOT', '')
    parent_folder = cfg.get('parent_folder', '')
    if not vce_root_env or not parent_folder:
        print("ERROR: VCE_ROOT env var and config.parent_folder must be set.", file=sys.stderr)
        sys.exit(2)

    docs_dir = Path(vce_root_env) / parent_folder / 'Documentation'
    output_dir = Path(os.environ.get('ETL_OUTPUT_DIR') or (Path(vce_root_env) / parent_folder / 'CSV'))
    output_dir.mkdir(parents=True, exist_ok=True)

    cfg['_docs_dir'] = docs_dir
    cfg['_output_dir'] = output_dir
    return cfg


# ============================================================================
# STYLE / HEADING HELPERS
# ============================================================================

def heading_level(style_name: Optional[str]) -> int:
    """Return heading level 1-6 (1 = Title/Heading 1), 0 for body text.
    Tolerates both 'Heading N' and 'VCAA Heading N' style names."""
    if not style_name:
        return 0
    s = style_name.lower()
    m = re.search(r'heading\s*(\d)', s)
    if m:
        return int(m.group(1))
    if s == 'title':
        return 1
    return 0


def is_body_style(style_name: Optional[str]) -> bool:
    if not style_name:
        return True
    if heading_level(style_name) > 0:
        return False
    s = style_name.lower()
    return any(kw in s for kw in ('body', 'normal', 'text', 'bullet', 'list', 'plain', 'student response'))


def clean(text: str) -> str:
    return re.sub(r'\s+', ' ', text or '').strip()


# ============================================================================
# CONTENT CLASSIFICATION
# ============================================================================

# First match wins. Ordered from most-specific to most-generic.
# Patterns tuned against VCAA PE + Accounting exam reports. Covers the most
# common intro phrasings examiners use to signal enhanced/limited commentary.
_STRATEGY_RULES: List[Tuple[re.Pattern, str]] = [
    # StudentResponse markers — check FIRST so "example of a high-scoring
    # response" goes to StudentResponse, not EnhancedStrategies.
    (re.compile(r'\bfollowing is (?:an |a )?(?:example|sample)', re.I), 'StudentResponse'),
    (re.compile(r'\bsample (?:student )?response\b', re.I), 'StudentResponse'),
    (re.compile(r'\bexample response\b', re.I), 'StudentResponse'),
    (re.compile(r'\bexample of (?:an? |the )?\w+(?:\s+\w+)?\s+response', re.I), 'StudentResponse'),

    # EnhancedStrategies
    (re.compile(r'\bresponses? that scored (?:highly|well)', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bhigher?[- ]?scoring\b', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bhigh[- ]?scoring (?:responses?|students?)', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bstronger responses?\b', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bbetter responses?\b', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bsuccessful responses?\b', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bstudents who (?:scored|did) well', re.I), 'EnhancedStrategies'),
    (re.compile(r'\bto be awarded full marks?\b', re.I), 'EnhancedStrategies'),

    # LimitedStrategies
    (re.compile(r'\bcommon (?:errors?|issues?|mistakes?|problems?)\b', re.I), 'LimitedStrategies'),
    (re.compile(r'\blower[- ]?scoring\b', re.I), 'LimitedStrategies'),
    (re.compile(r'\bweaker responses?\b', re.I), 'LimitedStrategies'),
    (re.compile(r'\blimited responses?\b', re.I), 'LimitedStrategies'),
    (re.compile(r'\bmany (?:responses|students) (?:incorrectly|failed|did not)', re.I), 'LimitedStrategies'),
    (re.compile(r'\bstudents (?:struggled|often|incorrectly|did not|failed to)\b', re.I), 'LimitedStrategies'),
    (re.compile(r'\bstudents found .{0,40}difficult\b', re.I), 'LimitedStrategies'),
]


def classify(text: str) -> str:
    for pattern, label in _STRATEGY_RULES:
        if pattern.search(text):
            return label
    return 'ExaminerCommentary'


# ============================================================================
# SECTION DISCOVERY
# ============================================================================

def find_sections(doc, configured: Dict[str, Dict]) -> List[Tuple[str, str, int, int]]:
    """Return (section_key, section_name, start_idx, end_idx) for each section.

    The overview range (top-of-doc through the first section start) is handled
    separately and not included here. For flat reports with no section
    headings (e.g. Accounting), the whole doc is treated as the single
    configured section.
    """
    paras = doc.paragraphs
    if not configured:
        return []

    hits: List[Tuple[str, str, int]] = []
    for i, p in enumerate(paras):
        if heading_level(p.style.name) != 2:
            continue
        text = p.text.strip()
        m = re.match(r'^Section\s+([A-Za-z0-9]+)\b', text, re.I)
        if not m:
            continue
        key = m.group(1).upper()
        if key in configured:
            hits.append((key, text, i))

    if not hits:
        if len(configured) == 1:
            only_key = next(iter(configured))
            name = configured[only_key].get('name', only_key)
            return [(only_key, name, 0, len(paras))]
        return []

    sections: List[Tuple[str, str, int, int]] = []
    for idx, (key, text, start) in enumerate(hits):
        end = hits[idx + 1][2] if idx + 1 < len(hits) else len(paras)
        configured_name = configured.get(key, {}).get('name')
        sections.append((key, configured_name or text, start, end))
    return sections


# ============================================================================
# OVERVIEW EXTRACTION
# ============================================================================

def extract_overview_rows(doc, cfg: Dict, first_section_idx: int, year: str) -> List[Dict]:
    """Emit overview rows from the report preamble (everything up to
    first_section_idx). Rows are grouped under the Heading-2 they fall under
    (e.g. 'General comments', 'Specific information')."""
    paras = doc.paragraphs
    subject_code = cfg.get('subject_code', '')
    base_er = f'VCE{subject_code}ER'

    rows: List[Dict] = []
    current_header = ''
    current_section_code = ''
    section_seq: Dict[str, int] = {}

    def _section_code_for(header: str) -> str:
        h = header.strip().lower()
        if 'general comment' in h:
            return base_er + 'GC'
        if 'specific information' in h:
            return base_er + 'SI'
        if 'administrative' in h:
            return base_er + 'AD'
        if not h:
            return base_er + 'GC'
        initials = ''.join(w[0] for w in re.findall(r'[A-Za-z]+', h))[:3].upper() or 'XX'
        return base_er + initials

    end = first_section_idx if first_section_idx > 0 else len(paras)

    for i in range(end):
        p = paras[i]
        text = clean(p.text)
        if not text:
            continue
        level = heading_level(p.style.name)
        if level == 1:
            continue  # document title
        if level == 2:
            current_header = text
            current_section_code = _section_code_for(text)
            section_seq.setdefault(current_section_code, 0)
            continue
        if not current_header:
            current_header = 'General comments'
            current_section_code = _section_code_for(current_header)
            section_seq.setdefault(current_section_code, 0)

        section_seq[current_section_code] += 1
        rows.append({
            'SubjectArea': cfg.get('subject_area', ''),
            'Subject': cfg.get('subject', ''),
            'SubjectStreamCode': '',
            'SubjectStreamName': '',
            'Band': 'Year 12',
            'AssessmentType': 'Final Exam',
            'AssessmentInformationDetails': 'Post Exam Report',
            'AssessmentYears': year,
            'UnitASCode': cfg.get('all_unit_as_codes', ''),
            'ExamReportCode': base_er,
            'SectionCode': current_section_code,
            'SectionHeader': current_header,
            'SectionHeaderDescription': '',
            'SectionHeaderContent': text,
            'ReportCoverage': 'All sections',
            'Sequence': section_seq[current_section_code],
        })

    return rows


# ============================================================================
# MULTIPLE-CHOICE EXTRACTION
# ============================================================================

def _extract_mc_table(doc) -> List[Dict]:
    """Return [{q, correct, comment}] from the first MCQ-style table whose
    column headers include 'Question' AND ('Correct answer' | 'Answer')."""
    candidates = []
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [clean(cell.text).lower() for cell in table.rows[0].cells]
        if any('question' in h for h in headers) and any('correct' in h or h == 'answer' for h in headers):
            candidates.append(table)

    if not candidates:
        return []

    table = candidates[0]
    headers = [clean(cell.text).lower() for cell in table.rows[0].cells]
    q_col = next((i for i, h in enumerate(headers) if 'question' in h), 0)
    ans_col = next((i for i, h in enumerate(headers) if 'correct' in h or h == 'answer'), None)
    comment_col = next((i for i, h in enumerate(headers) if 'comment' in h), None)

    rows = []
    for r in table.rows[1:]:
        cells = [clean(c.text) for c in r.cells]
        if q_col >= len(cells):
            continue
        q = cells[q_col]
        if not q or not re.match(r'^\d+[a-z]?$', q, re.I):
            continue
        rows.append({
            'q': q,
            'correct': cells[ans_col] if ans_col is not None and ans_col < len(cells) else '',
            'comment': cells[comment_col] if comment_col is not None and comment_col < len(cells) else '',
        })
    return rows


# ============================================================================
# WRITTEN-QUESTION EXTRACTION
# ============================================================================

_QUESTION_HEADING_RE = re.compile(
    r'^(?:Question\s+)?(\d+[a-z]?(?:\s*\([ivx]+\))?)(?:\s*[.\-\u2013:]|\s*$)',
    re.I,
)


def _is_question_heading(text: str, level: int) -> Optional[str]:
    if level != 3:
        return None
    m = _QUESTION_HEADING_RE.match(text.strip())
    if not m:
        return None
    return re.sub(r'\s+', '', m.group(1)).lower()


def _is_bullet_style(style_name: Optional[str]) -> bool:
    if not style_name:
        return False
    s = style_name.lower()
    return 'bullet' in s or 'list' in s


def _extract_written_questions(doc, start_idx: int, end_idx: int) -> List[Tuple[str, List[Tuple[str, bool]]]]:
    """Return [(question_label, [(text, is_bullet), ...])] for the range."""
    paras = doc.paragraphs
    result: List[Tuple[str, List[Tuple[str, bool]]]] = []
    current_label: Optional[str] = None
    current_body: List[Tuple[str, bool]] = []

    def _flush():
        if current_label is not None:
            result.append((current_label, current_body[:]))

    for i in range(start_idx, end_idx):
        p = paras[i]
        text = clean(p.text)
        level = heading_level(p.style.name)
        q_label = _is_question_heading(text, level)
        if q_label:
            _flush()
            current_label = q_label
            current_body = []
            continue
        if current_label is None:
            continue
        if not text:
            continue
        if level in (1, 2):
            _flush()
            current_label = None
            current_body = []
            continue
        if is_body_style(p.style.name) or level >= 3:
            current_body.append((text, _is_bullet_style(p.style.name)))
    _flush()
    return result


# ============================================================================
# STRATEGY ROW ASSEMBLY
# ============================================================================

def _strategy_base(cfg: Dict, section_key: str, year: str) -> Dict:
    sec = cfg.get('exam_sections', {}).get(section_key, {})
    return {
        'SubjectArea': cfg.get('subject_area', ''),
        'Subject': cfg.get('subject', ''),
        'SubjectStreamCode': '',
        'SubjectStreamName': '',
        'Band': 'Year 12',
        'AssessmentType': 'Final Exam',
        'AssessmentInformationDetails': 'Post Exam Report',
        'AssessmentYears': year,
        'UnitASCode': sec.get('unit_as_code', ''),
        'ExamReportCode': sec.get('exam_report_code', ''),
    }


def build_strategy_rows(doc, cfg: Dict, section_key: str, section_name: str,
                        start_idx: int, end_idx: int, year: str) -> List[Dict]:
    sec_cfg = cfg.get('exam_sections', {}).get(section_key, {})
    section_type = sec_cfg.get('section_type', 'written')
    er_code = sec_cfg.get('exam_report_code', '')
    base = _strategy_base(cfg, section_key, year)
    rows: List[Dict] = []

    if section_type == 'multiple_choice':
        for mc in _extract_mc_table(doc):
            row = dict(base)
            row.update({
                'EQCode': f"{er_code}Q{mc['q']}",
                'SectionCode': er_code,
                'StrategyType': 'ExaminerCommentary',
                'ExamQuestion': mc['q'],
                'ResponseSource': f"Correct answer: {mc['correct']}" if mc['correct'] else '',
                'Content': mc['comment'],
                'Example': '',
                'Sequence': 1,
            })
            rows.append(row)
        return rows

    for q_label, body in _extract_written_questions(doc, start_idx, end_idx):
        # Classify each paragraph, then group consecutive same-label chunks.
        # Special rule: when a colon-terminated paragraph ("Common issues
        # included:") is followed by bullets, the bullets inherit the intro's
        # classification until a non-bullet paragraph breaks the run.
        # Classification with inheritance:
        # - A paragraph with its own strong signal (Enhanced / Limited /
        #   StudentResponse) claims its label.
        # - If that paragraph ends with ":" it arms inheritance — subsequent
        #   paragraphs that don't carry their own strong signal (i.e. fall
        #   back to ExaminerCommentary) adopt the inherited label. This keeps
        #   "Common issues included:" + its bullets together, and the
        #   "The following is an example of..." intro together with the
        #   indented student-response body that follows.
        labelled: List[Tuple[str, str]] = []
        inherited: Optional[str] = None
        for text, _is_bullet in body:
            own_label = classify(text)
            if own_label != 'ExaminerCommentary':
                labelled.append((own_label, text))
                inherited = own_label if text.rstrip().endswith(':') else None
            elif inherited:
                labelled.append((inherited, text))
                # Inheritance stays armed until a new strong signal fires.
            else:
                labelled.append((own_label, text))

        grouped: List[Tuple[str, List[str]]] = []
        for label, para in labelled:
            if grouped and grouped[-1][0] == label:
                grouped[-1][1].append(para)
            else:
                grouped.append((label, [para]))

        seq = 0
        for label, paras_chunk in grouped:
            seq += 1
            joined = '\n'.join(paras_chunk)
            row = dict(base)
            row.update({
                'EQCode': f"{er_code}Q{q_label}",
                'SectionCode': er_code,
                'StrategyType': label,
                'ExamQuestion': q_label,
                'ResponseSource': '',
                'Content': joined if label != 'StudentResponse' else '',
                'Example': joined if label == 'StudentResponse' else '',
                'Sequence': seq,
            })
            rows.append(row)
    return rows


# ============================================================================
# CSV WRITE
# ============================================================================

OVERVIEW_FIELDNAMES = [
    'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
    'Band', 'AssessmentType', 'AssessmentInformationDetails', 'AssessmentYears',
    'UnitASCode', 'ExamReportCode', 'SectionCode', 'SectionHeader',
    'SectionHeaderDescription', 'SectionHeaderContent', 'ReportCoverage', 'Sequence',
]

STRATEGY_FIELDNAMES = [
    'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
    'Band', 'AssessmentType', 'AssessmentInformationDetails', 'AssessmentYears',
    'UnitASCode', 'ExamReportCode', 'EQCode', 'SectionCode',
    'StrategyType', 'ExamQuestion', 'ResponseSource', 'Content', 'Example', 'Sequence',
]


def _open_csv(path: Path):
    try:
        return open(path, 'w', newline='', encoding='utf-8')
    except PermissionError:
        print(
            f"\nERROR: Cannot write to {path}\n"
            f"  The file is locked by another process (typically Excel or OneDrive\n"
            f"  sync). Close the file and re-run the pipeline.",
            file=sys.stderr,
        )
        sys.exit(2)


def write_rows(path: Path, fieldnames: List[str], rows: Iterable[Dict]):
    with _open_csv(path) as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, '') for k in fieldnames})


# ============================================================================
# MAIN
# ============================================================================

def _iter_report_docs(cfg: Dict) -> List[Tuple[str, Path]]:
    """Return [(year, local_path)] for each exam_report configured."""
    from urllib.parse import unquote as _unquote
    docs_dir: Path = cfg['_docs_dir']
    entries = cfg.get('documents', {}).get('exam_reports') or {}

    out: List[Tuple[str, Path]] = []
    if isinstance(entries, dict):
        for year, value in sorted(entries.items()):
            name = str(value)
            if name.startswith(('http://', 'https://')):
                name = _unquote(name.rsplit('/', 1)[-1])
            p = docs_dir / name
            if p.is_file():
                out.append((str(year), p))
    elif isinstance(entries, str):
        name = entries
        if name.startswith(('http://', 'https://')):
            name = _unquote(name.rsplit('/', 1)[-1])
        p = docs_dir / name
        if p.is_file():
            out.append(('', p))
    return out


def main():
    cfg = load_config()
    print('=' * 70)
    print(f"VCE {cfg.get('subject', '?')} Exam Report Parser")
    print('=' * 70)

    report_docs = _iter_report_docs(cfg)
    if not report_docs:
        print("  No exam_reports documents resolved on disk — nothing to do.")
        return

    configured_sections = cfg.get('exam_sections') or {}
    if not configured_sections:
        print("  No exam_sections configured in the subject JSON — skipping.")
        return

    subject_slug = cfg.get('subject_slug') or cfg.get('subject', '').lower().replace(' ', '_')
    output_dir: Path = cfg['_output_dir']

    all_overview: List[Dict] = []
    all_strategies: List[Dict] = []

    for year, path in report_docs:
        print(f"\n  Processing {year} report: {path.name}")
        doc = Document(str(path))

        sections = find_sections(doc, configured_sections)
        if not sections:
            print(f"    WARNING: no matching sections for configured keys "
                  f"{list(configured_sections.keys())} — skipping {path.name}")
            continue

        first_start = sections[0][2] if sections else len(doc.paragraphs)
        overview_rows = extract_overview_rows(doc, cfg, first_start, year)
        all_overview.extend(overview_rows)
        print(f"    Overview rows: {len(overview_rows)}")

        for section_key, section_name, start_idx, end_idx in sections:
            strat_rows = build_strategy_rows(doc, cfg, section_key, section_name,
                                             start_idx, end_idx, year)
            all_strategies.extend(strat_rows)
            print(f"    Section {section_key} ({section_name}): "
                  f"{len(strat_rows)} strategy rows")

    overview_path = output_dir / f'vcaa_vce_sd_{subject_slug}_exam_report_overview.csv'
    strategies_path = output_dir / f'vcaa_vce_sd_{subject_slug}_exam_report_strategies.csv'

    write_rows(overview_path, OVERVIEW_FIELDNAMES, all_overview)
    write_rows(strategies_path, STRATEGY_FIELDNAMES, all_strategies)

    print(f"\n  Wrote {overview_path} ({len(all_overview)} rows)")
    print(f"  Wrote {strategies_path} ({len(all_strategies)} rows)")
    print('\n' + '=' * 70)
    print('COMPLETE')
    print('=' * 70)


if __name__ == '__main__':
    main()
