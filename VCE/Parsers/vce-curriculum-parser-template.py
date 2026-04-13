"""
Victorian Senior Years English Curriculum Parser
=================================================

Parses VCE English/EAL study design document into CSV format
for Neo4j database loading.

Features:
- Study Overview extraction (Scope, Aims, Levels of achievement)
- F-10 compatible structure (reuses existing database schema)
- Dual student type support (English and EAL students)
- Automatic ContentHeader generation via phrase mappings
- Nested bullet preservation (Key Skills)
- Smart sentence splitting
- Outcome table extraction (English/EAL two-column tables)
- Sequence numbering aligned with exam assessment model

CSV Outputs:
- curriculum-overview.csv: Study overview sections (Scope, Aims, Levels)
- curriculum.csv: Content descriptors with elaborations (sequenced)
- curriculum-outcomes.csv: Key Knowledge + Key Skills for both student types

==============================================================================
VERSION CONTROL
==============================================================================
Version | Date           | Changes
--------|----------------|-----------------------------------------------------
3.2     | Feb 07, 2026   | FILENAME ALIGNMENT + OUTCOME TABLE FIX:
        |                | - Output filenames aligned to curriculum- prefix:
        |                |   overview.csv → curriculum-overview.csv
        |                |   outcomes.csv → curriculum-outcomes.csv
        |                |   curriculum.csv unchanged (was curriculum_v2.csv in 3.0)
        |                | - Matches import patterns v2.1 and GitHub structure
        |                | - BUG FIX: Outcome table detection was case-sensitive
        |                |   ('English Students' vs 'English students' in doc)
        |                |   Now lowercases before comparison (matching v1.9)
3.1     | Feb 07, 2026   | RENAME TO ENGLISH-SPECIFIC PARSER:
        |                | - Renamed from victorian-senior-years-curriculum-v3_0.py
        |                |   to victorian-senior-years-english-curriculum-v3_1.py
        |                | - Updated docstring to reflect English/EAL scope
        |                | - No logic changes - code identical to v3.0
        |                | - Separated to allow new Maths parser to be built
        |                |   alongside without conflicting
3.0     | Feb 05, 2026   | OVERVIEW + SEQUENCED CURRICULUM:
        |                | - NEW: Study Overview CSV extraction
        |                |   - Scope of study (paragraphs)
        |                |   - Aims (header + bullets)
        |                |   - Levels of achievement (subheaders, paragraphs, bullets)
        |                | - NEW: Sequence field in curriculum (resets per AoS)
        |                | - NEW: ContentType field ("Descriptor" or "Elaboration")
        |                | - Aligns curriculum model with exam assessment structure
        |                | - OverviewCode format: VCEE12OV{Section}{Seq:02d}
2.0     | Feb 04, 2026   | UNITASCODE + CONTENTCODE + SORT ORDER:
        |                | - Added UnitASCode field to curriculum CSV
        |                |   (e.g., VCEEU1AS1, VCEEU3AS2) for linking to
        |                |   Framework of Ideas data
        |                | - ContentCode format changed to align with UnitASCode:
        |                |   OLD: VCEE11AS11RET01 (Year+AS+Unit+Acronym+Seq)
        |                |   NEW: VCEEU1AS1RET01 (U+AS+Acronym+Seq)
        |                | - Curriculum rows now sorted by Unit then Area of
        |                |   Study for consistent ordering in CSV view
1.9     | Jan 19, 2026   | OUTCOME DESCRIPTION FIX:
        |                | - Each Outcome now gets its own OutcomeDescription
        |                | - Pre-collects all outcome tables at init
        |                | - Uses index counter to get next table in sequence
        |                | - Previously all Outcomes used the first table's text
1.8     | Jan 19, 2026   | ASSESSMENT SECTION FIX:
        |                | - Added detection for Assessment headings (VCAA Heading 2)
        |                | - Detects both "Assessment" (Units 1-2) and 
        |                |   "School-based assessment" (Units 3-4)
        |                | - Parser now stops processing Key Knowledge/Skills
        |                |   when it hits any assessment section
        |                | - Prevents Assessment content from bleeding into
        |                |   outcomes data
1.7     | Jan 19, 2026   | SORT ORDER FIX:
        |                | - Fixed sort to group Key Knowledge before Key Skills
        |                | - Sort key now: (student, unit, aos, outcome, type, seq)
        |                | - Within each Outcome: all Key Knowledge rows first,
        |                |   then all Key Skills rows
        |                | - Prevents interleaving of Knowledge/Skills rows
1.6     | Jan 19, 2026   | CSV STRUCTURE FIX:
        |                | - Added SubjectArea as first field in outcomes CSV
        |                | - Changed EAL suffix from 'A' to 'EAL' (e.g., OKKEAL)
        |                | - Changed Type field to 'Key Knowledge'/'Key Skills'
        |                |   (both words capitalised)
        |                | - Output sorted: all English rows first (by Unit,
        |                |   AoS, Sequence), then all EAL rows (same sort)
        |                | - CSV field order now matches required format exactly
1.5     | Jan 18, 2026   | ENHANCED CODE FORMAT: ContentCode now includes BOTH
        |                | AoS number AND Unit number (VCEE11AS11RET01 = Year 11,
        |                | AoS 1, Unit 1). OutcomeCode format updated: OKK/OKS
        |                | (Key Knowledge/Skills), suffix E/A (English/EAL).
        |                | Changed 'Code' field to 'OutcomeCode' in CSV.
1.4     | Jan 18, 2026   | CRITICAL BUG FIX: ContentCode now uses Unit number
        |                | instead of Area of Study number. Codes now unique
        |                | per unit (VCEE11AS1RET01 for Unit 1, VCEE11AS2RET01
        |                | for Unit 2, etc.)
1.3     | Jan 18, 2026   | Capitalized outcome statements for consistency
1.2     | Jan 18, 2026   | Added standardized output filenames
1.1     | Jan 18, 2026   | BUG FIX: Added flag-based Area of Study detection
        |                | to prevent "End-of-year examination" false positives.
        |                | Uses expecting_aos_title flag to only process 
        |                | Heading 3 as AoS title when immediately following
        |                | "Area of Study X" (Heading 2).
1.0     | Jan 15, 2026   | Initial release - full VCE English parsing
==============================================================================

Current Version: 3.2
"""

import re
import csv
import os
import sys
import json
from docx import Document
from pathlib import Path
from typing import List, Dict, Tuple, Optional

# Ensure emoji/Unicode prints don't crash Windows cp1252 consoles
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    pass


# ============================================================================
# CONFIGURATION
# ============================================================================

SUBJECT_CONFIG = {
    'name': 'English',
    'subject_code': 'E',
    'subject_area': 'English',
    'level': 'VCE',
    
    # Year mapping for Units
    'year_mapping': {
        'Unit 1': 'Year 11',
        'Unit 2': 'Year 11', 
        'Unit 3': 'Year 12',
        'Unit 4': 'Year 12'
    },
    
    'curriculum_doc': '/mnt/user-data/uploads/2024EnglishEALSD__2_.docx',
    'output_dir': '/home/claude/senior_english_v3_output',
    
    # All UnitASCodes for overview linking
    'all_unit_as_codes': 'VCEEU1AS1,VCEEU1AS2,VCEEU2AS1,VCEEU2AS2,VCEEU3AS1,VCEEU3AS2,VCEEU4AS1,VCEEU4AS2',
    
    # Overview sections to extract
    'overview_sections': {
        'Scope of study': {
            'code_suffix': 'SC',
            'sequence': 1
        },
        'Aims': {
            'code_suffix': 'AI',
            'sequence': 2
        },
        'Levels of achievement': {
            'code_suffix': 'LA',
            'sequence': 3
        }
    },
    
    # Hardcoded ContentHeader mappings (phrase -> header)
    # Order matters - first match wins
    'content_header_mappings': [
        # Unit 1 AoS 1 - Reading and exploring texts
        ('engage in reading and viewing texts', 'Reading and viewing texts'),
        ('exploration of texts involves', 'Exploration of texts'),
        ('participation in discussions', 'Participation in discussions'),
        ('for this outcome', 'Outcome'),
        ('opportunities to practise', 'Practice'),
        
        # Unit 1 AoS 2 - Crafting texts
        ('engage with and develop an understanding of effective', 'Effective and cohesive writing'),
        ('read and engage imaginatively and critically with mentor texts', 'Mentor texts'),
        ('individual and shared reading of mentor texts', 'Shared reading'),
        ('employ and experiment with the qualities', 'Experimenting with writing'),
        ('mentor texts can include', 'Mentor text types'),
        ('negotiate ideas and mentor texts', 'Text selection'),
        
        # Unit 2 AoS 1 - Reading and exploring texts
        ('develop their reading and viewing skills', 'Reading and viewing skills'),
        ('read or view a text, engaging with the ideas', 'Engaging with ideas'),
        ('developing analytical writing', 'Analytical writing'),
        ('read and explore one set text', 'Set text'),
        
        # Unit 2 AoS 2 - Exploring argument
        ('explore how arguments are constructed', 'Argument construction'),
        ('suitable texts for study should reflect', 'Text selection'),
        ('practise analysing persuasive texts', 'Analysing persuasive texts'),
        ('craft their writing using evidence', 'Using evidence'),
        ('employ their understanding of argument', 'Creating point of view'),
        ('mentor texts provide opportunities', 'Mentor text opportunities'),
        
        # Unit 3 AoS 1 - Reading and responding to texts
        ('build on their understanding of the text', 'Building understanding'),
        ('read and study one set text', 'Set text'),
        ('analytical and critical understanding', 'Analytical understanding'),
        ('write analytically about a text', 'Analytical writing'),
        ('sustained analytical writing', 'Sustained analytical writing'),
        ('students study one text', 'Set text study'),
        ('students apply reading', 'Applying reading skills'),
        
        # Unit 3 AoS 2 - Creating texts
        ('develop skills in creating', 'Creating texts'),
        ('framework of ideas', 'Framework of Ideas'),
        ('mentor texts provide models', 'Mentor texts'),
        ('create their own texts', 'Creating own texts'),
        ('students work with mentor texts', 'Working with mentor texts'),
        ('students participate in collaborative', 'Collaborative learning'),
        ('students use and experiment with vocabulary', 'Vocabulary experimentation'),
        
        # Unit 4 AoS 1 - Reading and responding to texts  
        ('further develop their analytical', 'Analytical skills'),
        ('compare texts', 'Comparing texts'),
        
        # Unit 4 AoS 2 - Analysing argument
        ('develop their understanding of argument', 'Understanding argument'),
        ('analyse the use of argument', 'Analysing argument'),
        ('written and spoken language', 'Language analysis'),
        ('students must explore and analyse', 'Exploring and analysing'),
        ('students apply their understanding', 'Applying understanding'),
        
        # Generic fallbacks
        ('in this area of study', 'Overview'),
        ('students are expected', 'Expectations'),
        ('assessment', 'Assessment'),
    ]
}


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace while preserving line breaks."""
    if not text:
        return ""
    text = re.sub(r'[^\S\n]+', ' ', text)
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    return '\n'.join(lines)


def capitalize_first_letter(text: str) -> str:
    """
    Capitalize the first letter of text.
    
    Examples:
    - "inferential reading strategies" -> "Inferential reading strategies"
    - "the ways purpose..." -> "The ways purpose..."
    """
    if not text:
        return ""
    text = text.strip()
    if not text:
        return ""
    return text[0].upper() + text[1:]


def split_into_sentences(text: str) -> List[str]:
    """
    Smart sentence splitting that handles common abbreviations.
    """
    protected = text
    abbreviations = [
        (r'\bDr\.', 'Dr<DOT>'),
        (r'\bMr\.', 'Mr<DOT>'),
        (r'\bMrs\.', 'Mrs<DOT>'),
        (r'\bMs\.', 'Ms<DOT>'),
        (r'\bProf\.', 'Prof<DOT>'),
        (r'\be\.g\.', 'e<DOT>g<DOT>'),
        (r'\bi\.e\.', 'i<DOT>e<DOT>'),
        (r'\betc\.', 'etc<DOT>'),
        (r'\bvs\.', 'vs<DOT>'),
    ]
    
    for pattern, replacement in abbreviations:
        protected = re.sub(pattern, replacement, protected)
    
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', protected)
    sentences = [s.replace('<DOT>', '.') for s in sentences]
    sentences = [s.strip() for s in sentences if s.strip()]
    
    return sentences


def remove_bullet(text: str) -> str:
    """Remove bullet point marker from start of text."""
    text = text.strip()
    bullet_markers = ['•', '–', '-', '○', '▪', '·', '●', '◦']
    for marker in bullet_markers:
        if text.startswith(marker):
            return text[len(marker):].strip()
    return text


def is_bullet_point(text: str) -> bool:
    """Check if text starts with a bullet point marker."""
    text = text.strip()
    bullet_markers = ['•', '–', '-', '○', '▪', '·', '●', '◦']
    return any(text.startswith(marker) for marker in bullet_markers)


def generate_content_header(text: str, mappings: list) -> str:
    """
    Generate a content header using hardcoded phrase mappings.
    Falls back to extracting first few words if no match found.
    """
    if not text:
        return ""
    
    text_lower = text.lower()
    
    # Check mappings in order (first match wins)
    for phrase, header in mappings:
        if phrase in text_lower:
            return header
    
    # Fallback: extract first few meaningful words
    clean = re.sub(r'[^\w\s]', '', text)
    words = clean.split()
    skip_words = {'in', 'this', 'the', 'a', 'an', 'and', 'of', 'to', 'for', 'with', 'that', 'are', 'is'}
    header_words = [w for w in words[:8] if w.lower() not in skip_words][:4]
    
    return ' '.join(word.capitalize() for word in header_words) if header_words else "Content"


def generate_acronym(text: str) -> str:
    """
    Generate an acronym from area of study title.
    
    Examples:
    - "Reading and exploring texts" -> "RET"
    - "Crafting texts" -> "CT"
    - "Exploring argument" -> "EA"
    """
    if not text:
        return "XX"
    
    # Remove common words
    skip_words = {'and', 'the', 'of', 'to', 'a', 'an', 'in', 'on', 'for'}
    
    words = text.split()
    acronym_chars = []
    
    for word in words:
        clean_word = re.sub(r'[^\w]', '', word).lower()
        if clean_word and clean_word not in skip_words:
            acronym_chars.append(clean_word[0].upper())
    
    return ''.join(acronym_chars) if acronym_chars else "XX"


# ============================================================================
# OVERVIEW PARSER
# ============================================================================

class OverviewParser:
    """Extracts Study Overview sections from VCE study design documents."""
    
    def __init__(self, doc: Document, config: Dict):
        self.doc = doc
        self.config = config
        self.subject = config['name']
        self.subject_area = config['subject_area']
        self.overview_data = []
        
    def parse(self):
        """Extract overview sections."""
        print("\n" + "="*70)
        print("EXTRACTING STUDY OVERVIEW")
        print("="*70)
        
        self._extract_scope_of_study()
        self._extract_aims()
        self._extract_levels_of_achievement()
        
        print(f"\nTotal overview rows: {len(self.overview_data)}")
        
    def _find_section_start(self, heading_text: str) -> Optional[int]:
        """Find the paragraph index where a section starts."""
        for i, para in enumerate(self.doc.paragraphs):
            style = para.style.name if para.style else ''
            if 'Heading' in style and para.text.strip() == heading_text:
                return i
        return None
    
    def _generate_overview_code(self, section_suffix: str, seq: int) -> str:
        """Generate overview code like VCEE12OVSC01."""
        return f"VCE{self.config['subject_code']}12OV{section_suffix}{seq:02d}"
    
    def _add_row(self, section: str, section_suffix: str, content_type: str, 
                 description: str, sequence: int, subheader: str = '', url: str = '',
                 unit_as_code: str = None):
        """Add a row to overview data."""
        self.overview_data.append({
            'SubjectArea': self.subject_area,
            'Subject': self.subject,
            'Band': 'Years 11 and 12',
            'UnitASCode': unit_as_code if unit_as_code else self.config.get('all_unit_as_codes', ''),
            'OverviewCode': self._generate_overview_code(section_suffix, sequence),
            'Section': section,
            'SectionHeader': subheader,
            'ContentType': content_type,
            'Description': clean_text(description),
            'URL': url,
            'Sequence': sequence
        })
    
    def _extract_scope_of_study(self):
        """Extract Scope of study section."""
        print("\n  Extracting: Scope of study")
        
        start_idx = self._find_section_start('Scope of study')
        if start_idx is None:
            print("    WARNING: Scope of study not found")
            return
        
        section = 'Scope of study'
        suffix = self.config['overview_sections'][section]['code_suffix']
        seq = 1
        
        # Process paragraphs after heading until next Heading 2
        for para in self.doc.paragraphs[start_idx + 1:]:
            style = para.style.name if para.style else ''
            text = para.text.strip()
            
            if not text:
                continue
                
            # Stop at next heading
            if 'Heading 2' in style or 'Heading 1' in style:
                break
            
            if 'VCAA body' in style:
                self._add_row(section, suffix, 'Paragraph', text, seq)
                seq += 1
        
        print(f"    Extracted {seq - 1} rows")
    
    def _extract_aims(self):
        """Extract Aims section."""
        print("\n  Extracting: Aims")
        
        start_idx = self._find_section_start('Aims')
        if start_idx is None:
            print("    WARNING: Aims not found")
            return
        
        section = 'Aims'
        suffix = self.config['overview_sections'][section]['code_suffix']
        seq = 1
        
        # Process paragraphs after heading until next Heading 2
        for para in self.doc.paragraphs[start_idx + 1:]:
            style = para.style.name if para.style else ''
            text = para.text.strip()
            
            if not text:
                continue
                
            # Stop at next heading
            if 'Heading 2' in style or 'Heading 1' in style:
                break
            
            if 'VCAA body' in style:
                # This is the header "This study enables students to:"
                self._add_row(section, suffix, 'Header', text, seq)
                seq += 1
            elif 'VCAA bullet' in style:
                # Remove bullet marker and add
                bullet_text = remove_bullet(text) if is_bullet_point(text) else text
                self._add_row(section, suffix, 'Bullet', bullet_text, seq)
                seq += 1
        
        print(f"    Extracted {seq - 1} rows")
    
    def _extract_levels_of_achievement(self):
        """Extract Levels of achievement section."""
        print("\n  Extracting: Levels of achievement")
        
        start_idx = self._find_section_start('Levels of achievement')
        if start_idx is None:
            print("    WARNING: Levels of achievement not found")
            return
        
        section = 'Levels of achievement'
        suffix = self.config['overview_sections'][section]['code_suffix']
        seq = 1
        current_subheader = ''
        current_unit_as_code = None  # Track which unit group we're in
        
        # Unit-specific codes
        units_12_codes = 'VCEEU1AS1,VCEEU1AS2,VCEEU2AS1,VCEEU2AS2'
        units_34_codes = 'VCEEU3AS1,VCEEU3AS2,VCEEU4AS1,VCEEU4AS2'
        
        # VCE Administrative Handbook URL
        handbook_url = 'https://www.vcaa.vic.edu.au/administration/vce-vcal-handbook/Pages/index.aspx'
        
        # Process paragraphs after heading until next Heading 2 (Authentication)
        for para in self.doc.paragraphs[start_idx + 1:]:
            style = para.style.name if para.style else ''
            text = para.text.strip()
            
            if not text:
                continue
                
            # Stop at Authentication heading
            if 'Heading 2' in style and text == 'Authentication':
                break
            if 'Heading 1' in style:
                break
            
            if 'Heading 3' in style:
                # Subheader like "Units 1 and 2" or "Units 3 and 4"
                current_subheader = text
                if 'Units 1 and 2' in text:
                    current_unit_as_code = units_12_codes
                elif 'Units 3 and 4' in text:
                    current_unit_as_code = units_34_codes
                self._add_row(section, suffix, 'Subheader', text, seq, text, 
                              unit_as_code=current_unit_as_code)
                seq += 1
            elif 'VCAA body' in style:
                # Add URL if paragraph mentions VCE Administrative Handbook
                if 'VCE Administrative Handbook' in text:
                    self._add_row(section, suffix, 'Paragraph', text, seq, current_subheader, 
                                  url=handbook_url, unit_as_code=current_unit_as_code)
                else:
                    self._add_row(section, suffix, 'Paragraph', text, seq, current_subheader,
                                  unit_as_code=current_unit_as_code)
                seq += 1
                current_subheader = ''  # Only first para gets subheader
            elif 'VCAA bullet' in style:
                bullet_text = remove_bullet(text) if is_bullet_point(text) else text
                self._add_row(section, suffix, 'Bullet', bullet_text, seq,
                              unit_as_code=current_unit_as_code)
                seq += 1
        
        print(f"    Extracted {seq - 1} rows")


# ============================================================================
# CURRICULUM PARSER
# ============================================================================

class SeniorYearsParser:
    """
    Parser for VCE/Senior Years curriculum documents.
    """
    
    def __init__(self, config: Dict):
        self.config = config
        self.subject = config['name']
        self.subject_code = config['subject_code']
        self.subject_area = config['subject_area']
        self.year_mapping = config['year_mapping']
        
        # State tracking
        self.current_unit = None
        self.current_unit_description = ''
        self.current_year = None
        self.current_aos_num = None
        self.current_aos_title = None
        self.current_aos_full = None
        self.current_aos_acronym = None
        self.current_unit_as_code = None
        self.current_outcome = None
        self.current_outcome_description = {}  # Track per student type
        
        # Content tracking
        self.content_sequence = 1  # Reset per AoS
        self.content_code_sequence = 1  # For ContentCode
        self.current_content_description = None
        self.current_content_header = None
        self.current_content_code = None
        
        # Outcome tracking — keyed by student type from config (defaults to dual English+EAL)
        self._student_types = list(self.config.get('student_types') or (['English', 'EAL'] if self.config.get('has_dual_student_types') else [self.subject]))
        self.knowledge_sequence = {st: 0 for st in self._student_types}
        self.skill_sequence = {st: 0 for st in self._student_types}
        
        # Flags
        self.in_key_knowledge = False
        self.in_key_skills = False
        self.expecting_aos_title = False
        
        # Data storage
        self.curriculum_data = []
        self.outcomes_data = []
        
        # Pre-collect outcome tables
        self.outcome_tables = []
        self.outcome_table_index = 0
        
        self.doc = None
    
    
    def parse(self):
        """
        Parse the curriculum document.
        """
        print("="*70)
        print(f"Parsing VCE {self.subject} Study Design")
        print("="*70)
        
        doc_path = self.config['curriculum_doc']
        print(f"\nLoading: {doc_path}")
        
        self.doc = Document(doc_path)
        
        # Pre-collect outcome tables
        self._collect_outcome_tables()
        
        print(f"\nFound {len(self.doc.paragraphs)} paragraphs")
        print(f"Found {len(self.doc.tables)} tables")
        print(f"Found {len(self.outcome_tables)} outcome tables")
        
        self._parse_document()
        
        print(f"\n" + "-"*50)
        print(f"Curriculum content: {len(self.curriculum_data)} rows")
        print(f"Outcomes content: {len(self.outcomes_data)} rows")
    
    
    def _collect_outcome_tables(self):
        """
        Pre-collect all outcome tables from the document.
        These are tables with headers containing 'English Students' and 'EAL Students'.
        """
        for table in self.doc.tables:
            if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
                header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
                if 'english students' in header_text or 'eal students' in header_text:
                    self.outcome_tables.append(table)
    
    
    def _generate_subject_slug(self) -> str:
        """Generate slug for filenames.

        Prefers config['subject_slug'] when present (e.g. 'pe', 'history_revolutions').
        Otherwise falls back to <subject>_eal for legacy English+EAL dual-stream configs.
        """
        explicit = self.config.get('subject_slug')
        if explicit:
            return str(explicit).lower().replace('-', '_')
        base = self.subject.lower().replace(' ', '_')
        if self.config.get('has_dual_student_types'):
            return base + '_eal'
        return base
    
    
    def _generate_content_code(self) -> str:
        """
        Generate content code using UnitASCode pattern.
        Format: VCEEU{unit}AS{aos}{acronym}{seq:02d}
        Example: VCEEU1AS1RET01
        """
        unit_num = self.current_unit.replace('Unit ', '')
        return f"VCE{self.subject_code}U{unit_num}AS{self.current_aos_num}{self.current_aos_acronym}{self.content_code_sequence:02d}"
    
    
    def _generate_unit_as_code(self) -> str:
        """
        Generate UnitASCode for linking to Framework.
        Format: VCEEU{unit}AS{aos}
        Example: VCEEU1AS1
        """
        unit_num = self.current_unit.replace('Unit ', '')
        return f"VCE{self.subject_code}U{unit_num}AS{self.current_aos_num}"
    
    
    def _student_type_suffix(self, student_type: str) -> str:
        """Return OutcomeCode suffix for a given student type.

        For non-dual configs (e.g. PE), no suffix. For English+EAL dual:
        'E' for English, 'EAL' for EAL (preserved legacy behaviour).
        """
        if not self.config.get('has_dual_student_types'):
            return ''
        return 'E' if student_type == 'English' else 'EAL'

    def _current_outcome_num(self) -> str:
        """Extract numeric portion from current outcome heading (e.g. 'Outcome 1' -> '1')."""
        if not self.current_outcome:
            return '1'
        m = re.search(r'(\d+)', self.current_outcome)
        return m.group(1) if m else '1'

    def _generate_knowledge_code(self, student_type: str) -> str:
        """Generate Outcome Knowledge code.

        New format: VCE{subject_code}U{unit}AS{aos}O{outcome}KK{suffix?}
        Example: VCEPEU1AS1O1KK (PE, non-dual) or VCEEU1AS1O1KKE (English).
        """
        unit_num = self.current_unit.replace('Unit ', '')
        outcome_num = self._current_outcome_num()
        suffix = self._student_type_suffix(student_type)
        return f"VCE{self.subject_code}U{unit_num}AS{self.current_aos_num}O{outcome_num}KK{suffix}"

    def _generate_skill_code(self, student_type: str) -> str:
        """Generate Outcome Skill code.

        New format: VCE{subject_code}U{unit}AS{aos}O{outcome}KS{suffix?}
        """
        unit_num = self.current_unit.replace('Unit ', '')
        outcome_num = self._current_outcome_num()
        suffix = self._student_type_suffix(student_type)
        return f"VCE{self.subject_code}U{unit_num}AS{self.current_aos_num}O{outcome_num}KS{suffix}"
    
    
    def _parse_document(self):
        """Main parsing loop through document paragraphs."""
        
        for i, para in enumerate(self.doc.paragraphs):
            style = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect Unit headings (VCAA Heading 1)
            # Accept "Unit 1" (English) or "Unit 1: Title" (PE etc.)
            if 'VCAA Heading 1' in style:
                unit_match = re.match(r'^Unit (\d)(?:\s*[:\-\u2013]\s*(.+?))?\s*$', text)
                if unit_match:
                    unit_title = (unit_match.group(2) or '').strip()
                    self._handle_unit(f"Unit {unit_match.group(1)}", unit_title)
                    continue
            
            # Detect "Area of Study X" (VCAA Heading 2) - just the number
            if 'VCAA Heading 2' in style:
                aos_match = re.match(r'^Area of Study (\d)$', text)
                if aos_match:
                    self.current_aos_num = aos_match.group(1)
                    self.expecting_aos_title = True
                    continue
                
                # Detect Assessment section - stop processing outcomes
                if text in ['Assessment', 'School-based assessment']:
                    self.in_key_knowledge = False
                    self.in_key_skills = False
                    continue
            
            # Detect AoS title (VCAA Heading 3) - only if expecting it
            if 'VCAA Heading 3' in style and self.expecting_aos_title:
                self._handle_area_of_study_title(text)
                self.expecting_aos_title = False
                continue
            
            # Skip if we haven't established context
            if not self.current_unit or not self.current_aos_full:
                continue
            
            # Detect Outcome heading
            if text.startswith('Outcome ') and len(text) < 15:
                self._handle_outcome(text)
                continue
            
            # Detect Key Knowledge section
            if text == 'Key knowledge':
                self._start_key_knowledge()
                continue
            
            # Detect Key Skills section  
            if text == 'Key skills':
                self._start_key_skills()
                continue
            
            # Process content based on current state
            if self.in_key_knowledge:
                self._process_key_knowledge(para)
            elif self.in_key_skills:
                self._process_key_skills(para)
            elif ('VCAA body' in style or style == 'Plain Text') and self.current_aos_full:
                self._process_curriculum_content(para)
    
    
    def _handle_unit(self, text: str, unit_description: str = ''):
        """Handle Unit heading.

        text: normalised "Unit N" (lookup key for year_mapping).
        unit_description: title that followed "Unit N:" in the heading
            (e.g. "The human body in motion"). Empty for subjects like
            English where the heading is just "Unit 1".
        """
        self.current_unit = text
        self.current_unit_description = unit_description
        self.current_year = self.year_mapping.get(text, 'Year 11')
        self.current_aos_num = None
        self.current_aos_title = None
        self.current_aos_full = None
        self.in_key_knowledge = False
        self.in_key_skills = False
        label = f"{text}: {unit_description}" if unit_description else text
        print(f"\n  {label} ({self.current_year})")
    
    
    def _handle_area_of_study_title(self, title: str):
        """Handle Area of Study title."""
        self.current_aos_title = title
        self.current_aos_full = f"Area of Study {self.current_aos_num} - {title}"
        self.current_aos_acronym = generate_acronym(title)
        self.current_unit_as_code = self._generate_unit_as_code()
        
        # Reset sequences for new AoS
        self.content_sequence = 1
        self.content_code_sequence = 1
        self.current_content_description = None
        self.current_content_header = None
        self.current_content_code = None
        
        self.in_key_knowledge = False
        self.in_key_skills = False
        
        print(f"    Area of Study {self.current_aos_num}: {title} [{self.current_aos_acronym}]")
    
    
    def _handle_outcome(self, text: str):
        """Handle Outcome heading."""
        self.current_outcome = text
        self.in_key_knowledge = False
        self.in_key_skills = False
        
        # Reset outcome-related sequences
        self.knowledge_sequence = {st: 0 for st in self._student_types}
        self.skill_sequence = {st: 0 for st in self._student_types}
        
        # Get the outcome description from the next table
        self._get_outcome_descriptions()
    
    
    def _get_outcome_descriptions(self):
        """Get outcome descriptions from the next outcome table."""
        self.current_outcome_description = {st: '' for st in self._student_types}
        
        if self.outcome_table_index < len(self.outcome_tables):
            table = self.outcome_tables[self.outcome_table_index]
            self.outcome_table_index += 1
            
            # Find English and EAL descriptions from table
            if len(table.rows) > 1:
                for row in table.rows[1:]:
                    if len(row.cells) >= 2:
                        # Check which column has what
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if 'English students' in cell_text.lower() or cell_text.startswith('On completion'):
                                # This might be the description
                                pass
                
                # Simpler approach: assume first data row has both descriptions
                if len(table.rows) > 1 and len(table.rows[1].cells) >= 2:
                    self.current_outcome_description['English'] = table.rows[1].cells[0].text.strip()
                    self.current_outcome_description['EAL'] = table.rows[1].cells[1].text.strip()
    
    
    def _process_curriculum_content(self, para):
        """Process curriculum body content with sequencing."""
        text = para.text.strip()
        
        if not text:
            return
        
        # Stop processing if we hit assessment-related content or cross-unit content
        assessment_stops = [
            'award of satisfactory completion',
            'the areas of study',
            'all assessments at units',
            'for this unit students are required',
            'suitable tasks for assessment',
            'where teachers allow students to choose',
            'school-based assessment',
            'external assessment',
            'assessment of levels',
            'students are expected to read widely in units 3',  # Units 3-4 text selection
            'in units 3 and 4',
            'five texts across the units 3 and 4',
            'the text type selected for study',
            'only one of the selected texts',
            'no text studied at units 1',
            'either one of the texts selected',
            # Additional Units 3-4 assessment stops
            'level of achievement',
            'school-assessed coursework',
            'schoolassessed coursework',
            'where teachers provide a range',
            'types of range of forms',
            'achievement on unit',
            'for the achievement of unit',
            'written text constructed in consideration',
            'examination will contribute',
            'examination will be completed',
            'vcaa publishes specifications',
            'contribution to final assessment',
            'percentage contribution',
        ]
        text_lower = text.lower()
        for stop_phrase in assessment_stops:
            if stop_phrase in text_lower:
                return
        
        # Split into sentences
        sentences = split_into_sentences(text)
        
        if not sentences:
            return
        
        # First sentence is the ContentDescription
        first_sentence = sentences[0]
        
        # Check if this is a new content block or continuation
        # New block if first sentence is different from current
        if first_sentence != self.current_content_description:
            # New content block
            self.content_code_sequence += 1 if self.current_content_description else 0
            self.current_content_description = first_sentence
            self.current_content_header = generate_content_header(first_sentence, self.config.get('content_header_mappings', []))
            self.current_content_code = self._generate_content_code()
            
            # Add descriptor row
            self.curriculum_data.append({
                'SubjectArea': self.subject_area,
                'Subject': self.subject,
                'Band': self.current_year,
                'Unit': self.current_unit,
                'UnitDescription': self.current_unit_description,
                'UnitASCode': self.current_unit_as_code,
                'AreaofStudy': self.current_aos_full,
                'ContentCode': self.current_content_code,
                'ContentType': 'Descriptor',
                'ContentHeader': self.current_content_header,
                'ContentDescription': first_sentence,
                'Elaboration': '',
                'Sequence': self.content_sequence
            })
            self.content_sequence += 1
        
        # Remaining sentences are elaborations
        for sentence in sentences[1:]:
            self.curriculum_data.append({
                'SubjectArea': self.subject_area,
                'Subject': self.subject,
                'Band': self.current_year,
                'Unit': self.current_unit,
                'UnitDescription': self.current_unit_description,
                'UnitASCode': self.current_unit_as_code,
                'AreaofStudy': self.current_aos_full,
                'ContentCode': self.current_content_code,
                'ContentType': 'Elaboration',
                'ContentHeader': self.current_content_header,
                'ContentDescription': '',
                'Elaboration': sentence,
                'Sequence': self.content_sequence
            })
            self.content_sequence += 1
    
    
    def _start_key_knowledge(self):
        """Start Key Knowledge section."""
        self.in_key_knowledge = True
        self.in_key_skills = False
    
    
    def _start_key_skills(self):
        """Start Key Skills section."""
        self.in_key_knowledge = False
        self.in_key_skills = True
        self.current_skill_parent = {st: None for st in self._student_types}
    
    
    def _process_key_knowledge(self, para):
        """Process Key Knowledge statements."""
        text = para.text.strip()
        
        if not text or text in ['Key knowledge', 'Key skills']:
            return
        
        statement = remove_bullet(text) if is_bullet_point(text) else text
        
        if not statement:
            return
        
        # Add for each configured student type (English+EAL dual, or single for non-dual subjects)
        for student_type in self._student_types:
            self.knowledge_sequence[student_type] += 1
            self.outcomes_data.append({
                'SubjectArea': self.subject_area,
                'Subject': self.subject,
                'Band': self.current_year,
                'Unit': self.current_unit,
                'UnitDescription': self.current_unit_description,
                'UnitASCode': self.current_unit_as_code,
                'AreaofStudy': self.current_aos_full,
                'OutcomeCode': self._generate_knowledge_code(student_type),
                'Outcome': self.current_outcome,
                'OutcomeStudentType': f"{student_type} Students",
                'OutcomeDescription': self.current_outcome_description.get(student_type, ''),
                'Type': 'Key Knowledge',
                'Statement': capitalize_first_letter(clean_text(statement)),
                'ParentStatement': '',
                'Level': 1,
                'Sequence': self.knowledge_sequence[student_type]
            })
    
    
    def _process_key_skills(self, para):
        """Process Key Skills statements with nesting support."""
        text = para.text.strip()
        style = para.style.name if para.style else 'Normal'
        
        if not text or text in ['Key knowledge', 'Key skills']:
            return
        
        statement = remove_bullet(text) if is_bullet_point(text) else text
        
        if not statement:
            return
        
        # Detect level from style name
        is_level_2 = 'level 2' in style.lower()
        
        # Add for each configured student type
        for student_type in self._student_types:
            self.skill_sequence[student_type] += 1
            
            if is_level_2:
                parent = getattr(self, f'current_skill_parent_{student_type}', None)
                self.outcomes_data.append({
                    'SubjectArea': self.subject_area,
                    'Subject': self.subject,
                    'Band': self.current_year,
                    'Unit': self.current_unit,
                    'UnitDescription': self.current_unit_description,
                    'UnitASCode': self.current_unit_as_code,
                    'AreaofStudy': self.current_aos_full,
                    'OutcomeCode': self._generate_skill_code(student_type),
                    'Outcome': self.current_outcome,
                    'OutcomeStudentType': f"{student_type} Students",
                    'OutcomeDescription': self.current_outcome_description.get(student_type, ''),
                    'Type': 'Key Skills',
                    'Statement': capitalize_first_letter(clean_text(statement)),
                    'ParentStatement': parent if parent else '',
                    'Level': 2,
                    'Sequence': self.skill_sequence[student_type]
                })
            else:
                self.outcomes_data.append({
                    'SubjectArea': self.subject_area,
                    'Subject': self.subject,
                    'Band': self.current_year,
                    'Unit': self.current_unit,
                    'UnitDescription': self.current_unit_description,
                    'UnitASCode': self.current_unit_as_code,
                    'AreaofStudy': self.current_aos_full,
                    'OutcomeCode': self._generate_skill_code(student_type),
                    'Outcome': self.current_outcome,
                    'OutcomeStudentType': f"{student_type} Students",
                    'OutcomeDescription': self.current_outcome_description.get(student_type, ''),
                    'Type': 'Key Skills',
                    'Statement': capitalize_first_letter(clean_text(statement)),
                    'ParentStatement': '',
                    'Level': 1,
                    'Sequence': self.skill_sequence[student_type]
                })
                
                # Track as potential parent if ends with colon
                if statement.endswith(':'):
                    setattr(self, f'current_skill_parent_{student_type}', statement)
                else:
                    setattr(self, f'current_skill_parent_{student_type}', None)
    
    
    def write_csv_files(self, overview_data: List[Dict] = None):
        """
        Write all parsed data to CSV files.
        
        Filename pattern: vcaa_vce_sd_{subject}_{type}.csv
        """
        output_dir = Path(self.config['output_dir'])
        output_dir.mkdir(parents=True, exist_ok=True)
        
        subject_slug = self._generate_subject_slug()
        
        def _open_csv(path):
            try:
                return open(path, 'w', newline='', encoding='utf-8')
            except PermissionError:
                print(
                    f"\nERROR: Cannot write to {path}\n"
                    f"  The file is locked by another process — most likely open in\n"
                    f"  Excel or being synced by OneDrive. Close the file (and any\n"
                    f"  Excel window showing it) and re-run the pipeline.",
                    file=sys.stderr,
                )
                sys.exit(2)

        # Stream columns (empty for non-stream subjects; populated for streams
        # like History Revolutions that share a parent folder).
        def _with_streams(rows, include_stream_unit_as_code=False, include_stream_header=False):
            for row in rows:
                row.setdefault('SubjectStreamCode', '')
                row.setdefault('SubjectStreamName', '')
                if include_stream_unit_as_code:
                    row.setdefault('StreamUnitASCode', '')
                if include_stream_header:
                    row.setdefault('StreamHeader', '')
            return rows

        # Write overview
        if overview_data:
            overview_fieldnames = [
                'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
                'Band', 'UnitASCode', 'OverviewCode',
                'Section', 'SectionHeader', 'ContentType', 'Description', 'URL', 'Sequence'
            ]
            _with_streams(overview_data)

            overview_file = output_dir / f'vcaa_vce_sd_{subject_slug}_curriculum-overview.csv'
            with _open_csv(overview_file) as f:
                writer = csv.DictWriter(f, fieldnames=overview_fieldnames)
                writer.writeheader()
                writer.writerows(overview_data)
            print(f"\n✓ Wrote {overview_file}")
            print(f"  {len(overview_data)} rows")
        
        # Write curriculum content (sorted by Unit, then AoS, then Sequence)
        if self.curriculum_data:
            curriculum_fieldnames = [
                'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
                'Band', 'Unit', 'UnitDescription', 'UnitASCode', 'StreamUnitASCode',
                'AreaofStudy', 'ContentCode', 'ContentType', 'StreamHeader', 'ContentHeader',
                'ContentDescription', 'Elaboration', 'Sequence'
            ]
            _with_streams(self.curriculum_data, include_stream_unit_as_code=True, include_stream_header=True)

            def curriculum_sort_key(row):
                unit_num = int(row['Unit'].replace('Unit ', ''))
                aos_match = re.search(r'Area of Study (\d+)', row['AreaofStudy'])
                aos_num = int(aos_match.group(1)) if aos_match else 0
                return (unit_num, aos_num, row['Sequence'])

            sorted_curriculum = sorted(self.curriculum_data, key=curriculum_sort_key)

            curriculum_file = output_dir / f'vcaa_vce_sd_{subject_slug}_curriculum.csv'
            with _open_csv(curriculum_file) as f:
                writer = csv.DictWriter(f, fieldnames=curriculum_fieldnames)
                writer.writeheader()
                writer.writerows(sorted_curriculum)
            print(f"\n✓ Wrote {curriculum_file}")
            print(f"  {len(sorted_curriculum)} rows")
        
        # Write outcomes (sorted by Unit, then AreaofStudy)
        if self.outcomes_data:
            outcome_fieldnames = [
                'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
                'Band', 'Unit', 'UnitDescription', 'UnitASCode', 'StreamUnitASCode',
                'AreaofStudy', 'StreamHeader', 'OutcomeCode', 'Outcome', 'OutcomeStudentType',
                'OutcomeDescription', 'Type', 'Statement', 'ParentStatement', 'Level', 'Sequence'
            ]
            _with_streams(self.outcomes_data, include_stream_unit_as_code=True, include_stream_header=True)
            
            def sort_key(row):
                unit_num = int(row['Unit'].replace('Unit ', ''))
                aos_match = re.search(r'Area of Study (\d+)', row['AreaofStudy'])
                aos_num = int(aos_match.group(1)) if aos_match else 0
                student_order = 0 if 'English' in row['OutcomeStudentType'] else 1
                outcome_match = re.search(r'Outcome (\d+)', row['Outcome'])
                outcome_num = int(outcome_match.group(1)) if outcome_match else 0
                type_order = 0 if row['Type'] == 'Key Knowledge' else 1
                return (unit_num, aos_num, student_order, outcome_num, type_order, row['Sequence'])
            
            sorted_outcomes = sorted(self.outcomes_data, key=sort_key)
            
            outcomes_file = output_dir / f'vcaa_vce_sd_{subject_slug}_curriculum-outcomes.csv'
            with _open_csv(outcomes_file) as f:
                writer = csv.DictWriter(f, fieldnames=outcome_fieldnames)
                writer.writeheader()
                writer.writerows(sorted_outcomes)
            print(f"\n✓ Wrote {outcomes_file}")
            print(f"  {len(sorted_outcomes)} rows")


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def _load_config_from_argv() -> dict:
    """If argv[1] is a pipeline config JSON, merge it into SUBJECT_CONFIG.

    Maps pipeline config fields (pe.json-style) onto the hardcoded
    SUBJECT_CONFIG keys this parser expects. Resolves curriculum_doc and
    output_dir from VCE_ROOT + parent_folder.
    """
    cfg = dict(SUBJECT_CONFIG)  # start from hardcoded defaults

    if len(sys.argv) < 2:
        return cfg

    config_path = sys.argv[1]
    if not config_path or not Path(config_path).is_file():
        return cfg

    with open(config_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Field remapping: pe.json uses 'subject', parser wants 'name'
    cfg['name'] = data.get('subject') or cfg.get('name', '')
    cfg['subject_code'] = data.get('subject_code') or cfg.get('subject_code', '')
    cfg['subject_area'] = data.get('subject_area') or cfg.get('subject_area', '')
    cfg['subject_slug'] = data.get('subject_slug') or cfg.get('subject_slug', '')
    cfg['has_dual_student_types'] = bool(data.get('has_dual_student_types', False))
    cfg['level'] = 'VCE'

    if data.get('year_mapping'):
        cfg['year_mapping'] = data['year_mapping']
    if data.get('overview_sections'):
        cfg['overview_sections'] = data['overview_sections']
    if 'all_unit_as_codes' in data:
        cfg['all_unit_as_codes'] = data.get('all_unit_as_codes', '') or ''
    if 'content_header_mappings' in data:
        raw_chm = data.get('content_header_mappings') or []
        # Accept either [[phrase, header], ...] or [{"phrase":..,"header":..}]
        norm = []
        for item in raw_chm:
            if isinstance(item, (list, tuple)) and len(item) == 2:
                norm.append((item[0], item[1]))
            elif isinstance(item, dict) and 'phrase' in item and 'header' in item:
                norm.append((item['phrase'], item['header']))
        cfg['content_header_mappings'] = norm

    # Resolve curriculum_doc: VCE_ROOT / parent_folder / Documentation / <study_design filename>
    # The pipeline sets VCE_ROOT env override when it's a URL (after downloading).
    vce_root = os.environ.get('VCE_ROOT', '')
    parent_folder = data.get('parent_folder', '')
    docs = data.get('documents') or {}
    study_design = docs.get('study_design', '')

    if study_design and vce_root and parent_folder:
        # If study_design is a URL, the pipeline downloaded it to
        #   <override_root>/<parent_folder>/Documentation/<filename>
        # and set VCE_ROOT to override_root. So only the filename matters here.
        filename = study_design.rsplit('/', 1)[-1] if study_design.startswith(('http://', 'https://')) else study_design
        candidate = Path(vce_root) / parent_folder / 'Documentation' / filename
        cfg['curriculum_doc'] = str(candidate)

    # Resolve output_dir: VCE_ROOT_ORIGINAL / parent_folder / graphRAG_csv / CSV
    # Prefer ETL_OUTPUT_DIR env if provided; else derive from VCE_ROOT.
    output_override = os.environ.get('ETL_OUTPUT_DIR', '')
    if output_override:
        cfg['output_dir'] = output_override
    elif vce_root and parent_folder:
        cfg['output_dir'] = str(Path(vce_root) / parent_folder / 'graphRAG_csv' / 'CSV')

    # Ensure output dir exists
    out_dir = cfg.get('output_dir', '')
    if out_dir:
        Path(out_dir).mkdir(parents=True, exist_ok=True)

    print(f"[argv config] subject={cfg.get('name')} code={cfg.get('subject_code')}")
    print(f"[argv config] curriculum_doc={cfg.get('curriculum_doc')}")
    print(f"[argv config] output_dir={cfg.get('output_dir')}")
    return cfg


if __name__ == '__main__':
    # Merge pipeline config (argv[1]) into SUBJECT_CONFIG
    SUBJECT_CONFIG = _load_config_from_argv()

    # Load document once
    doc = Document(SUBJECT_CONFIG['curriculum_doc'])

    # Parse overview
    overview_parser = OverviewParser(doc, SUBJECT_CONFIG)
    overview_parser.parse()

    # Parse curriculum
    curriculum_parser = SeniorYearsParser(SUBJECT_CONFIG)
    curriculum_parser.parse()

    # Write all CSV files
    curriculum_parser.write_csv_files(overview_data=overview_parser.overview_data)
    
    print("\n" + "="*70)
    print("✓ PARSING COMPLETE!")
    print("="*70)
    print("\nFiles created:")
    print(f"  1. Curriculum-Overview CSV (Scope, Aims, Levels)")
    print(f"  2. Curriculum CSV (with Sequence and ContentType)")
    print(f"  3. Curriculum-Outcomes CSV")
