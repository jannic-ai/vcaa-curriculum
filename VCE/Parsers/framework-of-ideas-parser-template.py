"""
Framework of Ideas Parser
==========================

Extracts the complete Framework of Ideas content from VCE English study design.
Includes Overview, Key Ideas table, Textual Forms, Audience and Context, 
Purpose, and Mentor Texts sections.

Output CSV:
SubjectArea,Subject,Band,Unit,UnitASCode,AreaofStudy,FWKeyIdeaCode,FWKeyIdea,FWKeyIdeaHeader,FWKeyIdeaDescription,FWKIElaboration,Sequence

Features:
- Extracts all Framework of Ideas sections (not just the table)
- Generates unique codes: VCEE34FWKI## pattern
- Distinguishes between Description (prose) and Elaboration (bullet items)
- Sequence resets for each section/key idea group
- Sentence splitting with line breaks for multi-sentence descriptions

==============================================================================
VERSION CONTROL
==============================================================================
Version | Date           | Changes
--------|----------------|-----------------------------------------------------
2.2     | Feb 04, 2026   | UNITASCODE FIELD: Added UnitASCode column with comma-
        |                | delimited values linking Framework to curriculum
        |                | AreaOfStudy nodes (VCEEU3AS2,VCEEU4AS2). Positioned
        |                | between Unit and AreaofStudy for relationship mapping.
2.1     | Feb 04, 2026   | CAPITALISATION: Added capitalise_first() function to
        |                | ensure first character is uppercase for FWKeyIdeaHeader,
        |                | FWKeyIdeaDescription, and FWKIElaboration fields.
2.0     | Feb 04, 2026   | COMPLETE RESTRUCTURE:
        |                | - New column structure with FWKeyIdeaCode, FWKeyIdea
        |                |   (category), FWKeyIdeaHeader, FWKeyIdeaDescription,
        |                |   FWKIElaboration, Sequence
        |                | - Extracts all sections: Overview, Key Ideas table,
        |                |   Textual Forms, Audience/Context, Purpose, Mentor Texts
        |                | - Code generation: VCEE34FWKI## (sequential)
        |                | - Unit changed to "Unit 3 and Unit 4"
        |                | - AreaofStudy changed to "Framework of Ideas"
        |                | - Multi-sentence descriptions joined with line breaks
1.2     | Jan 18, 2026   | FILENAME CONVENTION: Added standardized filename
        |                | generation using pattern: vcaa_vce_sd_{subject}_fw_key_ideas.csv
1.1     | Jan 18, 2026   | SENTENCE-SPLIT UPDATE: Splits elaborations into
        |                | individual sentences for bullet point display.
        |                | Now generates one row per sentence (~17 rows).
1.0     | Jan 18, 2026   | Initial release - extracts Framework table with
        |                | full multi-paragraph elaboration text (4 rows)
==============================================================================

Current Version: 2.2
"""

import csv
import re
from docx import Document

# Configuration
DOC_PATH = '/mnt/user-data/uploads/2024EnglishEALSD__2_.docx'
OUTPUT_DIR = '/home/claude'

# Fixed metadata
METADATA = {
    'SubjectArea': 'English',
    'Subject': 'English',
    'Band': 'Year 12',
    'Unit': 'Unit 3 and Unit 4',
    'UnitASCode': 'VCEEU3AS2,VCEEU4AS2',
    'AreaofStudy': 'Framework of Ideas'
}

# Global code counter
code_counter = 0


def generate_subject_slug(subject):
    """
    Generate subject slug for filenames.
    
    Examples:
    - "English" -> "english_eal"
    - "Psychology" -> "psychology"
    """
    if subject.lower() == 'english':
        return 'english_eal'
    return subject.lower().replace(' ', '_').replace('-', '_')


def generate_code():
    """Generate next VCEE34FWKI## code."""
    global code_counter
    code_counter += 1
    return f"VCEE34FWKI{code_counter:02d}"


def capitalise_first(text):
    """
    Ensure the first character of text is uppercase.
    Handles empty strings and preserves rest of text.
    """
    if not text:
        return text
    return text[0].upper() + text[1:] if len(text) > 1 else text.upper()


def split_sentences(text):
    """
    Split text into sentences, handling abbreviations.
    Returns list of sentences.
    """
    # Handle abbreviations like Dr., Mrs., e.g., etc.
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s+(?=[A-Z])', text)
    return [s.strip() for s in sentences if s.strip()]


def join_with_linebreaks(text):
    """
    Split text into sentences and join with line breaks.
    Used for Description fields.
    """
    sentences = split_sentences(text)
    return '\n'.join(sentences)


def create_row(fw_key_idea, header, description='', elaboration='', sequence=1):
    """Create a single data row with all fields."""
    return {
        **METADATA,
        'FWKeyIdeaCode': generate_code(),
        'FWKeyIdea': fw_key_idea,
        'FWKeyIdeaHeader': capitalise_first(header),
        'FWKeyIdeaDescription': capitalise_first(description),
        'FWKIElaboration': capitalise_first(elaboration),
        'Sequence': sequence
    }


def find_framework_section(doc):
    """
    Find the Framework of Ideas section in the document.
    Returns the starting paragraph index.
    """
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Framework of Ideas' and 'Heading 2' in para.style.name:
            return i
    return None


def find_key_ideas_table(doc):
    """Find the Key Ideas table in the document."""
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            header_row = table.rows[0]
            col1 = header_row.cells[0].text.strip().lower()
            col2 = header_row.cells[1].text.strip().lower()
            
            if 'key idea' in col1 and 'elaboration' in col2:
                return table
    return None


def extract_overview(doc, start_idx):
    """
    Extract Overview section (paragraphs 340-341).
    Returns single row with combined description.
    """
    data = []
    
    # Get paragraphs after the Framework heading
    para_340 = doc.paragraphs[start_idx + 1].text.strip()
    para_341 = doc.paragraphs[start_idx + 2].text.strip()
    
    # Combine with line break
    description = f"{para_340}\n{para_341}"
    
    data.append(create_row(
        fw_key_idea='KeyIdea',
        header='Overview',
        description=description,
        sequence=1
    ))
    
    return data


def extract_key_ideas_table(doc):
    """
    Extract the 4 Key Ideas from the table.
    Each key idea gets a header row + elaboration rows.
    """
    data = []
    table = find_key_ideas_table(doc)
    
    if not table:
        print("WARNING: Key Ideas table not found!")
        return data
    
    # Process each key idea (rows 1-4, skipping header)
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        
        key_idea_name = row.cells[0].text.strip()
        elaboration_full = row.cells[1].text.strip()
        
        # Header row for this key idea (sequence 1, no elaboration)
        data.append(create_row(
            fw_key_idea='KeyIdea',
            header=key_idea_name,
            sequence=1
        ))
        
        # Split elaborations into sentences
        sentences = split_sentences(elaboration_full)
        
        # Elaboration rows (sequence 2+)
        for seq, sentence in enumerate(sentences, start=2):
            data.append(create_row(
                fw_key_idea='KeyIdea',
                header=key_idea_name,
                elaboration=sentence,
                sequence=seq
            ))
    
    return data


def extract_textual_forms(doc):
    """
    Extract Textual Forms section.
    Single row with description (sentences joined with line break).
    """
    data = []
    
    # Find the Textual forms heading
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Textual forms' and 'Heading 4' in para.style.name:
            # Next paragraph is the content
            content = doc.paragraphs[i + 1].text.strip()
            
            # Split sentences and join with line breaks
            description = join_with_linebreaks(content)
            
            data.append(create_row(
                fw_key_idea='TextualForms',
                header='Textual Forms',
                description=description,
                sequence=1
            ))
            break
    
    return data


def extract_audience_context(doc):
    """
    Extract Audience and Context section.
    Header+description row, then 6 bullet elaborations.
    """
    data = []
    
    # Find the Audience and context heading
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Audience and context' and 'Heading 4' in para.style.name:
            # Next paragraph is the intro description
            intro = doc.paragraphs[i + 1].text.strip()
            
            # Header row with description
            data.append(create_row(
                fw_key_idea='AudienceContext',
                header='Audience and context',
                description=intro,
                sequence=1
            ))
            
            # Next 6 paragraphs are bullet points (VCAA bullet style)
            seq = 2
            for j in range(i + 2, i + 8):
                if j < len(doc.paragraphs):
                    bullet_text = doc.paragraphs[j].text.strip()
                    if bullet_text and 'bullet' in doc.paragraphs[j].style.name.lower():
                        data.append(create_row(
                            fw_key_idea='AudienceContext',
                            header='Audience and context',
                            elaboration=bullet_text,
                            sequence=seq
                        ))
                        seq += 1
            break
    
    return data


def extract_purpose(doc):
    """
    Extract Purpose section.
    Header+description row (2 paragraphs combined), then 4 verb elaborations.
    """
    data = []
    
    # Find the Purpose heading
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Purpose' and 'Heading 4' in para.style.name:
            # Next two paragraphs are the intro description
            para_1 = doc.paragraphs[i + 1].text.strip()
            para_2 = doc.paragraphs[i + 2].text.strip()
            
            # Combine with line break
            description = f"{para_1}\n{para_2}"
            
            # Header row with description
            data.append(create_row(
                fw_key_idea='Purpose',
                header='Purpose',
                description=description,
                sequence=1
            ))
            
            # Next 4 paragraphs are the verb elaborations (Express, Explain, Reflect, Argue)
            seq = 2
            for j in range(i + 3, i + 7):
                if j < len(doc.paragraphs):
                    verb_text = doc.paragraphs[j].text.strip()
                    if verb_text:
                        data.append(create_row(
                            fw_key_idea='Purpose',
                            header='Purpose',
                            elaboration=verb_text,
                            sequence=seq
                        ))
                        seq += 1
            break
    
    return data


def extract_mentor_texts(doc):
    """
    Extract Mentor Texts section.
    Single row with description.
    """
    data = []
    
    # Find the Mentor texts heading
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Mentor texts' and 'Heading 4' in para.style.name:
            # Next paragraph is the content
            content = doc.paragraphs[i + 1].text.strip()
            
            data.append(create_row(
                fw_key_idea='MentorTexts',
                header='Mentor Texts',
                description=content,
                sequence=1
            ))
            break
    
    return data


def extract_framework_of_ideas():
    """Extract all Framework of Ideas content from the document."""
    global code_counter
    code_counter = 0  # Reset counter
    
    doc = Document(DOC_PATH)
    
    # Find Framework section start
    start_idx = find_framework_section(doc)
    if start_idx is None:
        print("ERROR: Framework of Ideas section not found!")
        return []
    
    print(f"Found Framework of Ideas at paragraph {start_idx}")
    
    # Extract all sections in order
    data = []
    
    print("\nExtracting Overview...")
    data.extend(extract_overview(doc, start_idx))
    
    print("Extracting Key Ideas table...")
    data.extend(extract_key_ideas_table(doc))
    
    print("Extracting Textual Forms...")
    data.extend(extract_textual_forms(doc))
    
    print("Extracting Audience and Context...")
    data.extend(extract_audience_context(doc))
    
    print("Extracting Purpose...")
    data.extend(extract_purpose(doc))
    
    print("Extracting Mentor Texts...")
    data.extend(extract_mentor_texts(doc))
    
    return data


def write_csv(data):
    """
    Write Framework of Ideas to CSV.
    
    Filename pattern: vcaa_vce_sd_{subject}_fw_key_ideas.csv
    Example: vcaa_vce_sd_english_eal_fw_key_ideas.csv
    """
    if not data:
        print("No data to write!")
        return None
    
    # Generate filename using subject slug
    subject_slug = generate_subject_slug(METADATA['Subject'])
    output_path = f"{OUTPUT_DIR}/vcaa_vce_sd_{subject_slug}_fw_key_ideas.csv"
    
    fieldnames = [
        'SubjectArea', 'Subject', 'Band', 'Unit', 'UnitASCode', 'AreaofStudy',
        'FWKeyIdeaCode', 'FWKeyIdea', 'FWKeyIdeaHeader', 
        'FWKeyIdeaDescription', 'FWKIElaboration', 'Sequence'
    ]
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(data)
    
    print(f"\n✓ Wrote {output_path}")
    print(f"  {len(data)} rows")
    
    return output_path


def print_summary(data):
    """Print summary of extracted data."""
    print("\n" + "="*70)
    print("EXTRACTION SUMMARY")
    print("="*70)
    
    # Count by FWKeyIdea type
    from collections import Counter
    type_counts = Counter(row['FWKeyIdea'] for row in data)
    
    print("\nRows by type:")
    for fw_type, count in type_counts.items():
        print(f"  {fw_type}: {count} rows")
    
    print(f"\nTotal rows: {len(data)}")
    print(f"Code range: VCEE34FWKI01 - VCEE34FWKI{len(data):02d}")


if __name__ == '__main__':
    print("="*70)
    print("FRAMEWORK OF IDEAS PARSER v2.0")
    print("="*70)
    
    data = extract_framework_of_ideas()
    
    if data:
        output_path = write_csv(data)
        print_summary(data)
        
        print("\n" + "="*70)
        print("✓ EXTRACTION COMPLETE!")
        print("="*70)
    else:
        print("\n✗ Extraction failed!")
