"""
Victorian Senior Years Exam Pre-Exam Information Parser
========================================================

Parses VCE English pre-exam assessment documents (Assessment Criteria, Expected Qualities)
into CSV format for Neo4j database loading.

Required Input Files:
- english-exam-assessment_critera-descriptors-24-27.docx

CSV Outputs:
- vcaa_vce_sd_english_final-exam-expected_qualities.csv (Grade descriptors per section)
- vcaa_vce_sd_english_final-exam-section_framework.csv (Assessment criteria per section)

==============================================================================
VERSION CONTROL
==============================================================================
Version | Date           | Changes
--------|----------------|-----------------------------------------------------
1.0     | Feb 04, 2026   | Initial release:
        |                | - Extracts Expected Qualities from Tables 0, 1, 2
        |                | - Extracts Assessment Criteria from paragraph bullets
        |                | - EQCode format: VCEE12EQA, VCEE12EQB, VCEE12EQC
        |                | - EQGradeCode format: VCEE12EQA10, VCEE12EQA09, etc.
        |                | - EQCriteriaCode format: VCEE12EQAAC01, etc.
        |                | - Expands mark ranges (9-10, 1-2) into individual rows
        |                | - Continuous sequence numbering within each section
        |                | - UnitASCode linkage to curriculum
        |                | - Row-per-bullet structure for granular queries
        |                | - Teacher-friendly EQSkill labels (A+ Student:, etc.)
==============================================================================

Current Version: 1.0
"""

import csv
from docx import Document
from pathlib import Path
from typing import List, Dict, Tuple


# ============================================================================
# CONFIGURATION
# ============================================================================

CONFIG = {
    'subject': 'English',
    'subject_area': 'English',
    'band': '12',
    'assessment_type': 'Final Exam',
    'assessment_info': {
        'details': 'Pre Exam Criteria',
        'years': '2024-2027'
    },
    
    # Input file
    'criteria_doc': '/mnt/user-data/uploads/english-exam-assessment_critera-descriptors-24-27.docx',
    
    # Output
    'output_dir': '/home/claude/vce_pre_exam_output',
    
    # Section configuration
    'sections': {
        'A': {
            'name': 'Analytical response to a text',
            'eq_code': 'VCEE12EQA',
            'unit_as_code': 'VCEEU3AS1,VCEEU4AS1',
            'table_index': 0
        },
        'B': {
            'name': 'Creating a text',
            'eq_code': 'VCEE12EQB',
            'unit_as_code': 'VCEEU3AS2,VCEEU4AS2',
            'table_index': 1
        },
        'C': {
            'name': 'Analysis of argument and language',
            'eq_code': 'VCEE12EQC',
            'unit_as_code': 'VCEEU3AS2,VCEEU4AS2',
            'table_index': 2
        }
    },
    
    # Mark to grade mapping (individual marks)
    'grade_mapping': {
        '10': ('A+', '10'),
        '9': ('A', '09'),
        '8': ('B+', '08'),
        '7': ('B', '07'),
        '6': ('C+', '06'),
        '5': ('C', '05'),
        '4': ('D+', '04'),
        '3': ('D', '03'),
        '2': ('E+', '02'),
        '1': ('E', '01'),
        '0': ('NR', '00')
    },
    
    # Ranges to expand (source mark text -> list of individual marks, highest first)
    'mark_ranges': {
        '9–10': ['10', '9'],
        '9-10': ['10', '9'],
        '1–2': ['2', '1'],
        '1-2': ['2', '1']
    },
    
    # Assessment criteria per section (from document paragraphs)
    'assessment_criteria': {
        'A': [
            'knowledge and understanding of the text, its structure, and the ideas, concerns and values it explores',
            'development of a coherent analysis in response to the topic',
            'use of evidence from the text to support the analysis',
            'use of fluent expression through appropriate use of vocabulary and conventions of Standard Australian English'
        ],
        'B': [
            'use of relevant idea(s) drawn from one Framework of Ideas, the title provided and at least one piece of stimulus material',
            'creation of a cohesive text that connects to a clear purpose(s) and incorporates an appropriate voice',
            'use of suitable text structure(s) and language features to create a text',
            'use of fluent expression, including the appropriate use of vocabulary'
        ],
        'C': [
            'understanding of contention, argument(s), and point of view',
            'analysis of the ways in which written and spoken language and visuals are used to present an argument(s) and to persuade an intended audience',
            'use of evidence from the text to support the analysis',
            'use of fluent expression through appropriate use of vocabulary and conventions of Standard Australian English'
        ]
    }
}


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def validate_input_file(config: Dict) -> Tuple[bool, str]:
    """Validate that required input file exists."""
    criteria_path = Path(config['criteria_doc'])
    
    if not criteria_path.exists():
        return False, f"ERROR: Required file missing: {config['criteria_doc']}"
    
    return True, "Input file present."


def get_grade_info(mark_text: str, config: Dict) -> Tuple[str, str]:
    """Convert mark text to grade and code suffix."""
    mark_text = mark_text.strip()
    
    if mark_text in config['grade_mapping']:
        return config['grade_mapping'][mark_text]
    
    # Handle unexpected formats
    print(f"  WARNING: Unknown mark format: '{mark_text}'")
    return ('?', '??')


def get_skill_label(grade: str) -> str:
    """Generate teacher-friendly skill label from grade."""
    return f"{grade} Student:"


def expand_mark_range(mark_text: str, config: Dict) -> List[str]:
    """Expand mark range text into individual marks (highest first)."""
    mark_text = mark_text.strip()
    
    # Check if it's a known range
    if mark_text in config['mark_ranges']:
        return config['mark_ranges'][mark_text]
    
    # Single mark
    return [mark_text]


# ============================================================================
# EXTRACTION FUNCTIONS
# ============================================================================

def extract_expected_qualities(doc: Document, config: Dict) -> List[List[str]]:
    """Extract Expected Qualities from all three section tables."""
    
    rows = []
    # Header row
    rows.append([
        'SubjectArea', 'Subject', 'Band', 'AssessmentType', 'AssessmentInformationDetails',
        'AssessmentYears', 'UnitASCode', 'EQCode', 'EQGradeCode', 'Mark', 'Grade',
        'EQSection', 'EQHeader', 'EQSkill', 'EQDescription', 'Sequence'
    ])
    
    for section_key in ['A', 'B', 'C']:
        section = config['sections'][section_key]
        table = doc.tables[section['table_index']]
        
        # Metadata for this section
        meta = [
            config['subject_area'],
            config['subject'],
            config['band'],
            config['assessment_type'],
            config['assessment_info']['details'],
            config['assessment_info']['years'],
            section['unit_as_code'],
            section['eq_code']
        ]
        
        seq = 1
        
        # Section header row
        header_text = f"Expected qualities for the mark range:\n\nSection {section_key} - {section['name']}"
        rows.append(meta + ['', '', '', section_key, header_text, '', '', str(seq)])
        seq += 1
        
        # Process each row in the table (skip header row)
        for table_row in table.rows[1:]:
            mark_text = table_row.cells[0].text.strip()
            qualities_text = table_row.cells[1].text.strip()
            
            # Split qualities into individual bullets
            bullets = [b.strip() for b in qualities_text.split('\n') if b.strip()]
            
            # Expand mark ranges (e.g., "9–10" becomes ["10", "9"])
            individual_marks = expand_mark_range(mark_text, config)
            
            # Create rows for each individual mark
            for mark in individual_marks:
                grade, code_suffix = get_grade_info(mark, config)
                eq_grade_code = f"{section['eq_code']}{code_suffix}"
                skill_label = get_skill_label(grade)
                
                first_bullet = True
                for bullet in bullets:
                    rows.append(meta + [
                        eq_grade_code,
                        mark,
                        grade,
                        section_key,
                        '',  # EQHeader empty for data rows
                        skill_label if first_bullet else '',  # Only on first bullet
                        bullet,
                        str(seq)
                    ])
                    seq += 1
                    first_bullet = False
        
        print(f"  Section {section_key}: {seq - 1} rows extracted")
    
    return rows


def extract_section_framework(config: Dict) -> List[List[str]]:
    """Extract Assessment Criteria for all three sections."""
    
    rows = []
    # Header row
    rows.append([
        'SubjectArea', 'Subject', 'Band', 'AssessmentType', 'AssessmentInformationDetails',
        'AssessmentYears', 'UnitASCode', 'EQCode', 'EQCriteriaCode', 'Section',
        'EQCriteriaSectionHeader', 'EQCriteriaType', 'EQCriteriaDescription', 'Sequence'
    ])
    
    for section_key in ['A', 'B', 'C']:
        section = config['sections'][section_key]
        criteria_list = config['assessment_criteria'][section_key]
        
        # Metadata for this section
        meta = [
            config['subject_area'],
            config['subject'],
            config['band'],
            config['assessment_type'],
            config['assessment_info']['details'],
            config['assessment_info']['years'],
            section['unit_as_code'],
            section['eq_code']
        ]
        
        for i, criteria in enumerate(criteria_list, 1):
            criteria_code = f"{section['eq_code']}AC{i:02d}"
            
            # Only include section header on first row (Sequence 1)
            section_header = section['name'] if i == 1 else ''
            
            rows.append(meta + [
                criteria_code,
                section_key,
                section_header,
                'Assessment Criteria',
                criteria,
                str(i)
            ])
        
        print(f"  Section {section_key}: {len(criteria_list)} criteria extracted")
    
    return rows


# ============================================================================
# MAIN EXECUTION
# ============================================================================

def write_csv(rows: List[List[str]], filepath: str):
    """Write rows to CSV file."""
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for row in rows:
            writer.writerow(row)
    print(f"  Created: {filepath} ({len(rows)-1} data rows)")


def main():
    """Main entry point."""
    print("=" * 70)
    print("VCE English Pre-Exam Information Parser v1.0")
    print("=" * 70)
    
    # Validate input file
    valid, msg = validate_input_file(CONFIG)
    if not valid:
        print(msg)
        return
    
    print(f"\n{msg}")
    
    # Create output directory
    output_dir = Path(CONFIG['output_dir'])
    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"\nOutput directory: {output_dir}")
    
    # Load document
    print("\nLoading document...")
    doc = Document(CONFIG['criteria_doc'])
    print(f"  Loaded: {CONFIG['criteria_doc']}")
    print(f"  Tables found: {len(doc.tables)}")
    
    # Extract Expected Qualities
    print("\n" + "-" * 50)
    print("Extracting Expected Qualities")
    print("-" * 50)
    eq_rows = extract_expected_qualities(doc, CONFIG)
    write_csv(eq_rows, output_dir / 'vcaa_vce_sd_english_final-exam-expected_qualities.csv')
    
    # Extract Section Framework (Assessment Criteria)
    print("\n" + "-" * 50)
    print("Extracting Section Framework (Assessment Criteria)")
    print("-" * 50)
    sf_rows = extract_section_framework(CONFIG)
    write_csv(sf_rows, output_dir / 'vcaa_vce_sd_english_final-exam-section_framework.csv')
    
    print("\n" + "=" * 70)
    print("EXTRACTION COMPLETE")
    print("=" * 70)
    print("\nFiles created:")
    print(f"  1. {output_dir / 'vcaa_vce_sd_english_final-exam-expected_qualities.csv'}")
    print(f"  2. {output_dir / 'vcaa_vce_sd_english_final-exam-section_framework.csv'}")


if __name__ == '__main__':
    main()
