#!/usr/bin/env python3
"""
VCE Outcome Examples Parser - CONFIG-DRIVEN
=============================================

Parses VCE sample teaching and learning activities documents into CSV format for Neo4j.
Links outcome examples to curriculum outcomes via UnitASCode.

Usage:
    python outcome-examples-parser-template.py <config.json>
    python outcome-examples-parser-template.py configs/english-language.json

CSV Output:
- vcaa_vce_sd_{slug}_outcome-examples.csv

Requires:
- The curriculum parser must have already run to generate the outcomes CSV
  (used to build UnitASCode -> AreaOfStudy lookup)
"""

import re
import csv
import json
import sys
import os
from docx import Document
from pathlib import Path
from typing import List, Dict, Optional


# ============================================================================
# PATH UTILITIES
# ============================================================================

def _long_path(p: str) -> str:
    """Prefix with \\\\?\\ on Windows to bypass MAX_PATH (260 char) limit."""
    if sys.platform == "win32" and not p.startswith("\\\\?\\"):
        return "\\\\?\\" + str(Path(p).resolve())
    return p


# ============================================================================
# TEXT UTILITIES
# ============================================================================

def fix_mojibake(text: str) -> str:
    """Fix double-encoded UTF-8 text (mojibake).

    Detects patterns like â€™ (should be ') where UTF-8 bytes were
    decoded as Windows-1252. Re-encodes to CP1252 bytes then decodes as UTF-8.
    """
    if not text or '\u00e2\u20ac' not in text:
        return text
    try:
        return text.encode('cp1252').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        return text


def normalize_unicode(text: str) -> str:
    """Normalize smart punctuation from DOCX to ASCII equivalents.

    IMPORTANT: All parsers must apply this to text extracted from documents.
    DOCX files contain smart quotes, em/en dashes, and other Unicode
    punctuation that displays as mojibake when tools read CSV files with
    wrong encoding. Normalizing to ASCII prevents this.
    """
    if not text:
        return ""
    text = fix_mojibake(text)
    replacements = {
        '\u2018': "'",   # left single quotation mark
        '\u2019': "'",   # right single quotation mark (apostrophe)
        '\u201C': '"',   # left double quotation mark
        '\u201D': '"',   # right double quotation mark
        '\u2013': '-',   # en dash
        '\u2014': '-',   # em dash
        '\u2026': '...',  # horizontal ellipsis
        '\u00A0': ' ',   # non-breaking space
        '\u2009': ' ',   # thin space
        '\u2003': ' ',   # em space
        '\u2002': ' ',   # en space
        '\u2008': ' ',   # punctuation space
        '\u00AD': '',    # soft hyphen
        '\u200B': '',    # zero-width space
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text


def sanitize_rows(rows):
    """Apply normalize_unicode to all string values in rows before CSV write.

    Handles both list-of-dicts and list-of-lists row formats.
    Call this immediately before writing to CSV as a safety net.
    """
    for row in rows:
        if isinstance(row, dict):
            for key, val in row.items():
                if isinstance(val, str):
                    row[key] = normalize_unicode(val)
        elif isinstance(row, list):
            for i, val in enumerate(row):
                if isinstance(val, str):
                    row[i] = normalize_unicode(val)
    return rows


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = normalize_unicode(text)
    text = text.replace('\t', ' ')
    text = re.sub(r' +', ' ', text)
    return text.strip()


def capitalize_first_letter(text: str) -> str:
    if not text:
        return ""
    text = text.strip()
    return text[0].upper() + text[1:] if text else ""


def extract_doc_url(doc: Document) -> str:
    """Extract URL from the top of the document (first hyperlink or URL text)."""
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    r_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

    for para in doc.paragraphs[:10]:
        # Check for hyperlink elements in XML
        for child in para._element:
            if child.tag == f'{w_ns}hyperlink':
                r_id = child.get(f'{r_ns}id')
                if r_id and r_id in doc.part.rels:
                    target = doc.part.rels[r_id].target_ref
                    if target and target.startswith('http'):
                        return target

        # Fallback: check for URL in plain text
        url_match = re.search(r'https?://\S+', para.text)
        if url_match:
            return url_match.group(0)

    return ''


# ============================================================================
# CONFIG LOADER
# ============================================================================

def load_config(config_path: str) -> Dict:
    """Load subject config from JSON file and resolve paths."""
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    # Resolve VCE root: prefer VCE_ROOT env var, fall back to path traversal
    vce_root_env = os.environ.get('VCE_ROOT', '')
    if vce_root_env:
        vce_dir = Path(vce_root_env)
    else:
        config_p = Path(config_path).resolve()
        etl_dir = config_p.parent.parent   # Up from configs/ to _etl/
        vce_dir = etl_dir.parent           # Up from _etl/ to VCE/

    subject_dir = vce_dir / config['parent_folder']
    docs_dir = subject_dir / "Documentation"

    # Resolve sample teaching document path from config.
    # Accept either a bare filename or a GitHub URL — the pipeline worker
    # downloads URL entries into docs_dir before we run, so we only ever
    # need the basename relative to docs_dir.
    def _filename_from(value: str) -> str:
        from urllib.parse import unquote
        if value.startswith(('http://', 'https://')):
            return unquote(value.rsplit('/', 1)[-1])
        return value

    sample_teaching = config.get('documents', {}).get('sample_teaching', '')
    if sample_teaching:
        config['examples_doc'] = str(docs_dir / _filename_from(sample_teaching))
    else:
        config['examples_doc'] = ''

    # ETL_OUTPUT_DIR env var (set by the pipeline worker) routes CSVs to the
    # shared <root>/<parent_folder>/CSV/ folder alongside the other subject
    # CSVs. Fallback to legacy graphRAG_csv/Curriculum layout for standalone
    # invocation.
    etl_output_override = os.environ.get('ETL_OUTPUT_DIR', '')
    config['output_dir'] = etl_output_override or str(subject_dir / "graphRAG_csv" / "Curriculum")

    # Resolve outcomes CSV path (generated by curriculum parser).
    # Preference order:
    #   1. ETL_OUTPUT_DIR — same run just produced it
    #   2. VCAA_REPO_DIR clone — curriculum parser was run in an earlier
    #      session and the CSV has already been published to GitHub
    #   3. Fallback to the configured output_dir (will fail cleanly if
    #      the file isn't there, matching the earlier error path)
    slug = config.get('subject_slug', config.get('subject', '').lower().replace(' ', '_'))
    outcomes_filename = f"vcaa_vce_sd_{slug}_curriculum-outcomes.csv"
    outcomes_candidates = [Path(config['output_dir']) / outcomes_filename]

    vcaa_repo = os.environ.get('VCAA_REPO_DIR', '')
    github_path = config.get('github_path', '')
    if vcaa_repo and github_path:
        outcomes_candidates.append(Path(vcaa_repo) / github_path / 'CSV' / outcomes_filename)

    resolved_outcomes = next((p for p in outcomes_candidates if p.is_file()), outcomes_candidates[0])
    config['outcomes_csv'] = str(resolved_outcomes)

    # Defaults
    config.setdefault('name', config.get('subject', 'UNKNOWN'))
    config.setdefault('subject_slug', config.get('subject', '').lower().replace(' ', '_'))

    return config


def load_outcomes_data(outcomes_csv_path: str) -> List[Dict]:
    """Load outcomes data from the curriculum parser's CSV output."""
    outcomes = []
    csv_path = Path(outcomes_csv_path)
    if not csv_path.exists():
        print(f"  WARNING: Outcomes CSV not found: {csv_path}")
        print("  (Run curriculum parser first to generate outcomes CSV)")
        return outcomes

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            outcomes.append(row)

    print(f"  Loaded {len(outcomes)} outcome rows from CSV")
    return outcomes


# ============================================================================
# SKIP MARKERS
# ============================================================================

SKIP_MARKERS = {
    'examples of learning activities',
    'detailed example',
}

# Skill verbs — if any of these appear in the first ~8 words of ExampleText,
# the row is classified as KS (Key Skills), otherwise KK (Key Knowledge).
SKILL_VERBS = {
    'identify', 'define', 'describe', 'apply', 'interpret',
    'discuss', 'compare', 'analyse', 'analyze', 'evaluate',
    'propose', 'justify',
}


def classify_kk_ks(text: str) -> str:
    """Return 'KS' if the opening words contain a skill verb, else 'KK'."""
    if not text:
        return 'KK'
    words = re.split(r'[\s,]+', text.lower())[:8]
    for w in words:
        if w in SKILL_VERBS:
            return 'KS'
    return 'KK'


def make_outcome_code(unit_as_code: str, outcome_label: str, text: str) -> str:
    """Build OutcomeCode like VCEBMU1AS1O1KK from UnitASCode + Outcome + text classification."""
    if not unit_as_code or not outcome_label:
        return ''
    m = re.search(r'(\d+)', outcome_label)
    if not m:
        return ''
    outcome_num = m.group(1)
    suffix = classify_kk_ks(text)
    return f"{unit_as_code}O{outcome_num}{suffix}"


# ============================================================================
# PARSER
# ============================================================================

def get_ilvl(para) -> Optional[int]:
    """Extract list indent level from paragraph's numPr XML element."""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    numPr = para._element.find(f'.//{ns}numPr')
    if numPr is None:
        return None
    ilvl_elem = numPr.find(f'{ns}ilvl')
    if ilvl_elem is not None:
        return int(ilvl_elem.get(f'{ns}val', '0'))
    return 0


def make_unit_as_code(subject_code: str, unit_num: int, outcome_num: Optional[int]) -> str:
    """Generate UnitASCode. If no outcome yet (general section), omit AS suffix."""
    if outcome_num is not None:
        return f"VCE{subject_code}U{unit_num}AS{outcome_num}"
    return f"VCE{subject_code}U{unit_num}"


def build_unit_outcomes(outcomes_data: List[Dict]) -> Dict[str, List[tuple]]:
    """Build ordered outcome sequence per unit from the curriculum outcomes CSV.

    Returns: {"Unit 1": [(UnitASCode, AreaofStudy, OutcomeLabel, UnitDescription), ...], ...}
    The list is in the order outcomes appear in the CSV (AoS 1 first, etc.).
    """
    unit_outcomes = {}
    for row in outcomes_data:
        unit = row.get('Unit', '')
        code = row.get('UnitASCode', '')
        aos = row.get('AreaofStudy', '')
        outcome = row.get('Outcome', '')
        unit_description = row.get('UnitDescription', '')
        if not unit or not code:
            continue
        if unit not in unit_outcomes:
            unit_outcomes[unit] = []
        # Only add first occurrence per UnitASCode (avoid duplicates from KK/KS rows)
        if not any(entry[0] == code for entry in unit_outcomes[unit]):
            unit_outcomes[unit].append((code, aos, outcome, unit_description))
    return unit_outcomes


def _unit_description_for(unit_outcomes: Dict[str, List[tuple]], unit_label: str) -> str:
    """Return the UnitDescription for a Unit label (e.g. 'Unit 1')."""
    entries = unit_outcomes.get(unit_label) or []
    for entry in entries:
        if len(entry) >= 4 and entry[3]:
            return entry[3]
    return ''


def parse_outcome_examples(doc: Document, config: Dict, unit_outcomes: Dict[str, List[tuple]]) -> List[Dict]:
    """Parse the sample teaching & learning activities document.

    Uses unit_outcomes (from curriculum CSV) as the source of truth for
    UnitASCode, AreaofStudy, and Outcome labels. The examples document
    is processed in order — each Outcome marker maps to the next entry
    in the unit's outcome sequence.
    """
    subject_area = config['subject_area']
    subject = config['name']
    subject_stream_code = config.get('subject_stream_code', '')
    subject_stream_name = config.get('subject_stream_name', '')
    subject_code = config['subject_code']
    year_mapping = config['year_mapping']

    examples_data = []
    current_unit_num = None
    current_unit_as_code = None
    current_outcome_label = None
    current_aos = ''
    current_unit_description = ''
    outcome_list = []           # Ordered outcomes for current unit
    outcome_index = 0           # Position in outcome_list
    in_detailed_example = False
    found_detailed_title = False
    expecting_outcome_desc = False
    sequence = 0
    last_parent_text = ''

    def _structural_row(content_type: str, text_val: str) -> Dict:
        """Emit a structural (non-activity) row using current parser state."""
        nonlocal sequence
        sequence += 1
        unit_label = f"Unit {current_unit_num}" if current_unit_num else ''
        band = year_mapping.get(unit_label, '') if unit_label else ''
        return {
            'SubjectArea': subject_area,
            'Subject': subject,
            'SubjectStreamCode': subject_stream_code,
            'SubjectStreamName': subject_stream_name,
            'Band': band,
            'Unit': unit_label,
            'UnitDescription': current_unit_description,
            'UnitASCode': current_unit_as_code or '',
            'AreaofStudy': current_aos,
            'Outcome': current_outcome_label or '',
            'OutcomeCode': '',
            'ExampleType': '',
            'ContentType': content_type,
            'ExampleText': capitalize_first_letter(clean_text(text_val)),
            'ParentText': '',
            'URLTitle': '',
            'URL': '',
            'Sequence': sequence,
        }

    for para in doc.paragraphs:
        style = para.style.name if para.style else 'Normal'
        text = para.text.strip()

        if not text:
            continue

        # --- Unit detection: Heading 3 style OR text-based (Normal) ---
        unit_match = re.match(r'^Unit\s+(\d+)', text)
        if unit_match and ('Heading 3' in style or style in ('Normal', 'Body Text')):
            current_unit_num = int(unit_match.group(1))
            unit_label = f"Unit {current_unit_num}"
            outcome_list = unit_outcomes.get(unit_label, [])
            outcome_index = 0
            current_unit_as_code = make_unit_as_code(subject_code, current_unit_num, None)
            current_outcome_label = None
            current_aos = ''
            current_unit_description = _unit_description_for(unit_outcomes, unit_label)
            in_detailed_example = False
            found_detailed_title = False
            expecting_outcome_desc = False
            sequence = 0
            last_parent_text = ''
            print(f"\n  Unit {current_unit_num} ({len(outcome_list)} outcomes from CSV)")
            examples_data.append(_structural_row('UnitHeading', text))
            continue

        # --- Outcome detection: Heading 4 style OR text-based (Normal) ---
        # Maps to the next outcome in sequence from the curriculum CSV,
        # regardless of the outcome number in the document text.
        outcome_match = re.match(r'^Outcome\s+(\d+)\s*:?\s*$', text)
        if outcome_match and ('Heading 4' in style or style in ('Normal', 'Body Text')):
            if current_unit_num is not None:
                if outcome_index < len(outcome_list):
                    # unit_outcomes entries are 4-tuples (code, aos, label,
                    # unit_description); older configs may still emit 3-tuples.
                    entry = outcome_list[outcome_index]
                    code, aos, label = entry[0], entry[1], entry[2]
                    current_unit_as_code = code
                    current_aos = aos
                    current_outcome_label = label
                    outcome_index += 1
                    print(f"    {label} -> {code} ({aos[:40]})")
                else:
                    # Fallback: more outcomes in doc than in CSV
                    doc_num = int(outcome_match.group(1))
                    current_unit_as_code = make_unit_as_code(subject_code, current_unit_num, doc_num)
                    current_outcome_label = f"Outcome {doc_num}"
                    current_aos = ''
                    print(f"    WARNING: Outcome {doc_num} exceeds CSV outcomes for Unit {current_unit_num}")
                in_detailed_example = False
                found_detailed_title = False
                expecting_outcome_desc = True
                sequence = 0
                last_parent_text = ''
                examples_data.append(_structural_row('OutcomeHeading', text))
            continue

        if current_unit_num is None:
            continue

        # --- Section markers (emit structural rows instead of skipping) ---
        text_lower = text.lower().strip()
        if text_lower == 'examples of learning activities':
            in_detailed_example = False
            found_detailed_title = False
            examples_data.append(_structural_row('SectionHeading', text))
            continue
        if text_lower.startswith('detailed example'):
            in_detailed_example = True
            found_detailed_title = False
            last_parent_text = ''
            continue

        # --- Outcome description (first Normal after Outcome marker) ---
        if expecting_outcome_desc:
            expecting_outcome_desc = False
            examples_data.append(_structural_row('OutcomeDescription', text))
            continue

        # --- Determine list indent level ---
        ilvl = get_ilvl(para)

        # --- Determine ExampleType ---
        if in_detailed_example:
            if not found_detailed_title:
                example_type = 'DetailedExample'
                found_detailed_title = True
            else:
                example_type = 'DetailedExampleStep'
        elif current_outcome_label is None:
            example_type = 'GeneralActivity'
        else:
            example_type = 'Activity'

        # --- ContentType: Bullet (list item) vs Paragraph (non-list) ---
        content_type = 'Bullet' if ilvl is not None else 'Paragraph'

        # --- Handle sub-items (ilvl>=1) ---
        parent_text = ''
        if ilvl is not None and ilvl >= 1:
            parent_text = last_parent_text
        elif ilvl == 0 or ilvl is None:
            last_parent_text = text

        sequence += 1
        unit_label = f"Unit {current_unit_num}"
        band = year_mapping.get(unit_label, 'Year 11')

        cleaned_text = capitalize_first_letter(clean_text(text))
        examples_data.append({
            'SubjectArea': subject_area,
            'Subject': subject,
            'SubjectStreamCode': subject_stream_code,
            'SubjectStreamName': subject_stream_name,
            'Band': band,
            'Unit': unit_label,
            'UnitDescription': current_unit_description,
            'UnitASCode': current_unit_as_code,
            'AreaofStudy': current_aos,
            'Outcome': current_outcome_label or '',
            'OutcomeCode': make_outcome_code(current_unit_as_code or '', current_outcome_label or '', cleaned_text),
            'ExampleType': example_type,
            'ContentType': content_type,
            'ExampleText': cleaned_text,
            'ParentText': clean_text(parent_text) if parent_text else '',
            'URLTitle': '',
            'URL': '',
            'Sequence': sequence
        })

    return examples_data


def write_csv(examples_data: List[Dict], config: Dict, doc_url: str = ''):
    """Write outcome examples CSV with header row prepended."""
    if not examples_data:
        print("  No outcome examples to write")
        return

    output_dir = Path(_long_path(config['output_dir']))
    output_dir.mkdir(parents=True, exist_ok=True)
    slug = config.get('subject_slug', 'unknown')

    fieldnames = [
        'SubjectArea', 'Subject', 'SubjectStreamCode', 'SubjectStreamName',
        'Band', 'Unit', 'UnitDescription', 'UnitASCode', 'AreaofStudy',
        'Outcome', 'OutcomeCode', 'ExampleType', 'ContentType', 'ExampleText',
        'ParentText', 'URLTitle', 'URL', 'Sequence'
    ]

    # Sort by unit, outcome, sequence
    def sort_key(row):
        unit_num = int(row['Unit'].replace('Unit ', ''))
        outcome_match = re.search(r'Outcome (\d+)', row['Outcome'])
        outcome_num = int(outcome_match.group(1)) if outcome_match else 0
        return (unit_num, outcome_num, row['Sequence'])

    sorted_data = sorted(examples_data, key=sort_key)

    # Prepend header row
    header_row = {
        'SubjectArea': config['subject_area'],
        'Subject': config['name'],
        'SubjectStreamCode': config.get('subject_stream_code', ''),
        'SubjectStreamName': config.get('subject_stream_name', ''),
        'Band': '',
        'Unit': '',
        'UnitDescription': '',
        'UnitASCode': config.get('all_unit_as_codes', ''),
        'AreaofStudy': '',
        'Outcome': '',
        'OutcomeCode': '',
        'ExampleType': '',
        'ContentType': 'Header',
        'ExampleText': '',
        'ParentText': f"{config['subject_area']} - {config['name']}",
        'URLTitle': f"{config['name']} - Teaching and Learning",
        'URL': doc_url,
        'Sequence': 0
    }
    sorted_data.insert(0, header_row)
    sanitize_rows(sorted_data)

    examples_file = output_dir / f'vcaa_vce_sd_{slug}_outcome-examples.csv'
    with open(_long_path(str(examples_file)), 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(sorted_data)
    print(f"\n  Wrote {examples_file.name} ({len(sorted_data)} rows, including header)")


# ============================================================================
# MAIN
# ============================================================================

def main():
    if len(sys.argv) < 2:
        print("Usage: python outcome-examples-parser-template.py <config.json>")
        sys.exit(1)

    config = load_config(sys.argv[1])

    print("\n" + "=" * 70)
    print("EXTRACTING OUTCOME EXAMPLES (Sample Teaching & Learning Activities)")
    print("=" * 70)
    print(f"\n  Subject: {config['name']}")

    # Check for sample teaching document
    if not config.get('examples_doc'):
        print("  No 'sample_teaching' document configured - skipping")
        sys.exit(0)

    doc_path = Path(config['examples_doc'])
    if not doc_path.exists():
        print(f"  ERROR: Document not found: {doc_path}")
        sys.exit(1)

    print(f"  Document: {doc_path.name}")

    # Load outcomes data from previously generated CSV
    outcomes_csv = Path(config['outcomes_csv'])
    if not outcomes_csv.exists():
        print(f"\n  ERROR: Curriculum outcomes CSV not found:")
        print(f"    {outcomes_csv}")
        print(f"\n  The curriculum parser must run first to generate this file.")
        print(f"  Run:  python curriculum-parser-template.py configs/{config.get('subject_slug', 'subject')}.json")
        sys.exit(1)

    outcomes_data = load_outcomes_data(config['outcomes_csv'])

    # Build ordered outcome sequence per unit from curriculum CSV
    unit_outcomes = build_unit_outcomes(outcomes_data)
    total_outcomes = sum(len(v) for v in unit_outcomes.values())
    if total_outcomes == 0:
        print(f"\n  ERROR: No outcomes found in curriculum CSV:")
        print(f"    {outcomes_csv}")
        print(f"\n  The file exists but contains no valid outcome rows.")
        print(f"  Re-run the curriculum parser to regenerate it.")
        sys.exit(1)
    print(f"  Unit outcomes: {len(unit_outcomes)} units, {total_outcomes} outcomes total")

    # Load and parse document
    print(f"\n  Loading: {doc_path.name}")
    doc = Document(str(doc_path))
    print(f"  Found {len(doc.paragraphs)} paragraphs")

    # Extract URL from top of document
    doc_url = extract_doc_url(doc)
    if doc_url:
        print(f"  Document URL: {doc_url}")
    else:
        print("  WARNING: No URL found at top of document")

    examples_data = parse_outcome_examples(doc, config, unit_outcomes)

    print(f"\n  Total outcome example rows: {len(examples_data)}")

    # Summary by UnitASCode
    code_counts = {}
    for row in examples_data:
        c = row['UnitASCode']
        code_counts[c] = code_counts.get(c, 0) + 1
    for code in sorted(code_counts):
        print(f"    {code}: {code_counts[code]} rows")

    # Summary by ExampleType
    type_counts = {}
    for row in examples_data:
        t = row['ExampleType']
        type_counts[t] = type_counts.get(t, 0) + 1
    for t in sorted(type_counts):
        print(f"    {t}: {type_counts[t]}")

    # Summary by OutcomeCode suffix (KK vs KS)
    kk_count = sum(1 for r in examples_data if r.get('OutcomeCode', '').endswith('KK'))
    ks_count = sum(1 for r in examples_data if r.get('OutcomeCode', '').endswith('KS'))
    no_code = sum(1 for r in examples_data if not r.get('OutcomeCode', ''))
    print(f"\n  OutcomeCode: {kk_count} KK, {ks_count} KS, {no_code} structural (no code)")

    write_csv(examples_data, config, doc_url)


if __name__ == '__main__':
    main()
